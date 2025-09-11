from PyQt5.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QPushButton, QLabel, 
    QTextEdit, QComboBox, QFileDialog, QMessageBox, QProgressBar,
    QGroupBox, QCheckBox, QTableWidget, QTableWidgetItem, QHeaderView,
    QTabWidget, QSplitter, QFrame, QScrollArea, QApplication
)
from PyQt5.QtCore import Qt, QThread, pyqtSignal, QTimer
from PyQt5.QtGui import QFont, QIcon
import os
import json
import hashlib
from datetime import datetime
from models.db_manager import DatabaseManager

# Word document generation (python-docx)
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    Document = None
    WD_PARAGRAPH_ALIGNMENT = None
    DOCX_AVAILABLE = False
    print("Cảnh báo: Chưa cài đặt thư viện python-docx. Cài đặt bằng: pip install python-docx")

# Sửa import statement này
from views.pages.report_ui.report_ui import Ui_Form

class ReportGenerator(QThread):
    """Background thread for generating comprehensive reports"""
    progress_updated = pyqtSignal(int, str)
    report_generated = pyqtSignal(str, str)  # file_path, format
    
    def __init__(self, case_id, report_type, options):
        super().__init__()
        self.case_id = case_id
        self.report_type = report_type
        self.options = options
        self.db = DatabaseManager()
        
    def run(self):
        try:
            # Check if DOCX library is available
            if not DOCX_AVAILABLE:
                raise Exception("Thiếu thư viện python-docx. Cài đặt bằng: pip install python-docx")

            self.db.connect()
            
            if self.report_type == "comprehensive":
                file_path = self.generate_comprehensive_report()
            elif self.report_type == "executive":
                file_path = self.generate_executive_summary()
            elif self.report_type == "technical":
                file_path = self.generate_technical_report()
            else:
                file_path = self.generate_chain_of_custody()
                
            if file_path:
                # Calculate hash and save to database
                with open(file_path, 'rb') as f:
                    content = f.read()
                    sha256 = hashlib.sha256(content).hexdigest()
                
                self.db.create_report(
                    case_id=self.case_id,
                    file_path=file_path,
                    format="DOCX",
                    sha256=sha256
                )
                
                self.report_generated.emit(file_path, self.report_type)
            else:
                self.report_generated.emit("", "error")
                
        except Exception as e:
            print(f"Report generation error: {e}")
            self.report_generated.emit("", "error")
        finally:
            self.db.disconnect()
    
    def generate_comprehensive_report(self):
        """Generate comprehensive case report with all details"""
        case_info = self.db.get_case_with_investigator(self.case_id)
        if not case_info or not isinstance(case_info, dict):
            print(f"ERROR - Invalid case info returned: {case_info}")
            return None
            
        # Get all case data with proper error handling
        try:
            artifacts = self.db.get_artifacts_by_case(self.case_id)
            if artifacts is None:
                artifacts = []
            elif not isinstance(artifacts, (list, tuple)):
                print(f"WARNING - Artifacts returned unexpected type: {type(artifacts)}, value: {artifacts}")
                artifacts = []
        except Exception as e:
            print(f"ERROR getting artifacts: {e}")
            artifacts = []
            
        try:
            results = self.db.get_results_by_case(self.case_id)
            if results is None:
                results = []
            elif not isinstance(results, (list, tuple)):
                print(f"WARNING - Results returned unexpected type: {type(results)}, value: {results}")
                results = []
        except Exception as e:
            print(f"ERROR getting results: {e}")
            results = []
            
        try:
            activity_logs = self.db.get_activity_logs(case_id=self.case_id)
            if activity_logs is None:
                activity_logs = []
            elif not isinstance(activity_logs, (list, tuple)):
                print(f"WARNING - Activity logs returned unexpected type: {type(activity_logs)}, value: {activity_logs}")
                activity_logs = []
        except Exception as e:
            print(f"ERROR getting activity logs: {e}")
            activity_logs = []

        # Create comprehensive Word document - use simplified version for now
        doc = self._create_simplified_docx(case_info, artifacts, results, activity_logs, self.options)
        
        # Save report
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"Case_{self.case_id}_Comprehensive_Report_{timestamp}.docx"

        archive_path = case_info.get('archive_path')
        if not archive_path:
            print(f"ERROR - No archive path found in case info: {case_info}")
            return None

        file_path = os.path.join(archive_path, filename)
        
        doc.save(file_path)
            
        return file_path
    
    def generate_executive_summary(self):
        """Generate executive summary report"""
        case_info = self.db.get_case_with_investigator(self.case_id)
        if not case_info or not isinstance(case_info, dict):
            print(f"ERROR - Invalid case info returned: {case_info}")
            return None
            
        # Get summary data
        artifacts = self.db.get_artifacts_by_case(self.case_id)
        results = self.db.get_results_by_case(self.case_id)
        
        # Ensure data is iterable
        if not isinstance(artifacts, (list, tuple)):
            artifacts = []
        if not isinstance(results, (list, tuple)):
            results = []
        
        doc = self._create_executive_docx(case_info, artifacts, results)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"Case_{self.case_id}_Executive_Summary_{timestamp}.docx"

        archive_path = case_info.get('archive_path')
        if not archive_path:
            print(f"ERROR - No archive path found in case info: {case_info}")
            return None

        file_path = os.path.join(archive_path, filename)
        
        doc.save(file_path)
            
        return file_path
    
    def generate_technical_report(self):
        """Generate technical detailed report"""
        case_info = self.db.get_case_with_investigator(self.case_id)
        if not case_info or not isinstance(case_info, dict):
            print(f"ERROR - Invalid case info returned: {case_info}")
            return None
            
        artifacts = self.db.get_artifacts_by_case(self.case_id)
        results = self.db.get_results_by_case(self.case_id)
        activity_logs = self.db.get_activity_logs(case_id=self.case_id)
        
        # Ensure data is iterable
        if not isinstance(artifacts, (list, tuple)):
            artifacts = []
        if not isinstance(results, (list, tuple)):
            results = []
        if not isinstance(activity_logs, (list, tuple)):
            activity_logs = []

        doc = self._create_technical_docx(case_info, artifacts, results, activity_logs)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"Case_{self.case_id}_Technical_Report_{timestamp}.docx"

        archive_path = case_info.get('archive_path')
        if not archive_path:
            print(f"ERROR - No archive path found in case info: {case_info}")
            return None

        file_path = os.path.join(archive_path, filename)
        
        doc.save(file_path)
            
        return file_path
    
    def generate_chain_of_custody(self):
        """Generate Chain of Custody report"""
        case_info = self.db.get_case_with_investigator(self.case_id)
        if not case_info or not isinstance(case_info, dict):
            print(f"ERROR - Invalid case info returned: {case_info}")
            return None
            
        artifacts = self.db.get_artifacts_by_case(self.case_id)
        activity_logs = self.db.get_activity_logs(case_id=self.case_id)
        
        # Ensure data is iterable
        if not isinstance(artifacts, (list, tuple)):
            artifacts = []
        if not isinstance(activity_logs, (list, tuple)):
            activity_logs = []

        doc = self._create_coc_docx(case_info, artifacts, activity_logs)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"Case_{self.case_id}_Chain_of_Custody_{timestamp}.docx"

        archive_path = case_info.get('archive_path')
        if not archive_path:
            print(f"ERROR - No archive path found in case info: {case_info}")
            return None

        file_path = os.path.join(archive_path, filename)
        
        doc.save(file_path)
            
        return file_path

    def _create_comprehensive_docx(self, case_info, artifacts, results, activity_logs):
        """Create comprehensive Word document"""
        from docx.shared import Pt, RGBColor
        from docx.enum.style import WD_STYLE_TYPE

        doc = Document()

        # Setup styles
        self._setup_word_styles(doc)

        # Header
        title = doc.add_paragraph("BÁO CÁO ĐIỀU TRA PHÁP Y SỐ TỔNG HỢP", style='CustomHeading1')
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        subtitle = doc.add_paragraph("CHUỖI BẢO QUẢN - TÍNH TOÀN VẸN BẰNG CHỨNG", style='CustomHeading2')
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        case_title = doc.add_paragraph(f"Mã số vụ án: {case_info['case_id']} - {case_info['title']}")
        case_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph("")

        # Case Information Section
        heading = doc.add_paragraph("THÔNG TIN VỤ ÁN", style='CustomHeading2')

        info_table = doc.add_table(rows=7, cols=2)  # 1 header + 6 data rows
        info_table.style = 'Table Grid'

        # Headers
        hdr_cells = info_table.rows[0].cells
        hdr_cells[0].text = "Thông tin"
        hdr_cells[1].text = "Giá trị"

        # Data rows
        rows_data = [
            ("Mã số vụ án", case_info['case_id']),
            ("Tên vụ án", case_info['title']),
            ("Trạng thái", case_info.get('status', 'Không xác định')),
            ("Điều tra viên", case_info.get('full_name', 'Không xác định')),
            ("Ngày tạo", case_info.get('created_at', 'Không xác định')),
            ("Đường dẫn lưu trữ", case_info.get('archive_path', 'Không xác định'))
        ]

        for i, (label, value) in enumerate(rows_data, 1):
            row_cells = info_table.rows[i].cells
            row_cells[0].text = str(label)
            row_cells[1].text = str(value)

        doc.add_page_break()

        # Evidence Section
        heading = doc.add_paragraph("BẰNG CHỨNG SỐ", style='CustomHeading2')

        if artifacts:
            evidence_table = doc.add_table(rows=1, cols=6)
            evidence_table.style = 'Table Grid'

            # Header row
            hdr_cells = evidence_table.rows[0].cells
            hdr_cells[0].text = "Mã số"
            hdr_cells[1].text = "Tên bằng chứng"
            hdr_cells[2].text = "Loại bằng chứng"
            hdr_cells[3].text = "Kích thước"
            hdr_cells[4].text = "Ngày thu thập"
            hdr_cells[5].text = "Mã băm SHA-256"

            # Data rows
            for artifact in artifacts:
                row_cells = evidence_table.add_row().cells
                row_cells[0].text = str(artifact['artefact_id'])
                row_cells[1].text = artifact['name']
                row_cells[2].text = artifact.get('evidence_type', 'Không xác định')
                row_cells[3].text = f"{artifact.get('size', 0):,} bytes"
                row_cells[4].text = artifact.get('collected_at', 'Không xác định')
                
                # Get SHA-256 hash from artifacts table or hashes table
                sha256_value = artifact.get('sha256', 'Không có')
                row_cells[5].text = str(sha256_value) if sha256_value is not None else 'Không có'
        else:
            doc.add_paragraph("Chưa có bằng chứng số nào được thu thập.")

        doc.add_page_break()

        # Analysis Results Section
        heading = doc.add_paragraph("KẾT QUẢ PHÂN TÍCH", style='CustomHeading2')

        if results:
            results_table = doc.add_table(rows=1, cols=4)
            results_table.style = 'Table Grid'

            # Header row
            hdr_cells = results_table.rows[0].cells
            hdr_cells[0].text = "Mã số"
            hdr_cells[1].text = "Công cụ sử dụng"
            hdr_cells[2].text = "Thời gian thực hiện"
            hdr_cells[3].text = "Tóm tắt kết quả"

            # Data rows
            for result in results:
                row_cells = results_table.add_row().cells
                row_cells[0].text = str(result['result_id'])
                row_cells[1].text = result.get('tool_used', 'Không xác định')
                row_cells[2].text = result.get('run_at', 'Không xác định')
                row_cells[3].text = result.get('summary', 'Không có thông tin')
        else:
            doc.add_paragraph("Chưa có kết quả phân tích nào được thực hiện.")

        doc.add_page_break()

        # Activity Log Section
        heading = doc.add_paragraph("NHẬT KÝ HOẠT ĐỘNG", style='CustomHeading2')

        if activity_logs:
            activity_table = doc.add_table(rows=1, cols=5)
            activity_table.style = 'Table Grid'

            # Header row
            hdr_cells = activity_table.rows[0].cells
            hdr_cells[0].text = "Thời gian"
            hdr_cells[1].text = "Hành động"
            hdr_cells[2].text = "Người thực hiện"
            hdr_cells[3].text = "Công cụ sử dụng"
            hdr_cells[4].text = "Chi tiết"

            # Data rows
            for log in activity_logs:
                row_cells = activity_table.add_row().cells
                row_cells[0].text = log.get('timestamp', 'Không xác định')
                row_cells[1].text = log.get('action', 'Không xác định')
                row_cells[2].text = log.get('username', 'Không xác định')
                row_cells[3].text = log.get('tool_used', 'Không có')
                row_cells[4].text = log.get('details', 'Không có')
        else:
            doc.add_paragraph("Chưa có hoạt động nào được ghi nhận trong hệ thống.")

        doc.add_page_break()

        # Chain of Custody Section
        heading = doc.add_paragraph("CHUỖI BẢO QUẢN BẰNG CHỨNG", style='CustomHeading2')

        coc_paragraph = doc.add_paragraph()
        coc_paragraph.add_run("Báo cáo này đảm bảo tính toàn vẹn của bằng chứng số:").bold = True

        coc_items = [
            "Tất cả bằng chứng đều có mã băm SHA-256 để xác minh tính toàn vẹn",
            "Nhật ký hoạt động ghi lại mọi thao tác với bằng chứng",
            "Dấu thời gian cho mọi hoạt động thu thập và phân tích",
            "Thông tin người thực hiện và công cụ sử dụng"
        ]

        for item in coc_items:
            p = doc.add_paragraph(item, style='List Bullet')

        doc.add_paragraph("")
        doc.add_paragraph(f"Báo cáo được tạo vào: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

        doc.add_page_break()

        # Recommendations Section
        heading = doc.add_paragraph("TÓM TẮT VÀ KHUYẾN NGHỊ", style='CustomHeading2')

        rec_paragraph = doc.add_paragraph()
        rec_paragraph.add_run("Dựa trên kết quả phân tích, vụ án này cần:").bold = True

        recommendations = [
            "Tiếp tục thu thập thêm bằng chứng nếu cần thiết",
            "Hoàn thiện và duy trì chuỗi bảo quản bằng chứng",
            "Chuẩn bị báo cáo chi tiết cho cơ quan có thẩm quyền",
            "Lưu trữ an toàn và bảo mật tất cả bằng chứng số"
        ]

        for rec in recommendations:
            p = doc.add_paragraph(rec, style='List Bullet')

        return doc

    def _create_executive_docx(self, case_info, artifacts, results):
        """Create executive summary Word document"""
        doc = Document()

        # Setup styles
        self._setup_word_styles(doc)

        # Header
        title = doc.add_paragraph("TÓM TẮT BAN GIÁM ĐỐC", style='CustomHeading1')
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        case_title = doc.add_paragraph(f"Vụ án {case_info['case_id']}: {case_info['title']}")
        case_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph("")

        # Case Summary Section
        heading = doc.add_paragraph("TÓM TẮT VỤ ÁN", style='CustomHeading2')

        summary_table = doc.add_table(rows=3, cols=2)
        summary_table.style = 'Table Grid'

        # Data rows
        rows_data = [
            ("Điều tra viên", case_info.get('full_name', 'Không xác định')),
            ("Ngày tạo", case_info.get('created_at', 'Không xác định')),
            ("Trạng thái", case_info.get('status', 'Không xác định'))
        ]

        for i, (label, value) in enumerate(rows_data):
            row_cells = summary_table.rows[i].cells
            row_cells[0].text = label
            row_cells[1].text = value

        doc.add_paragraph("")

        # Statistics Section
        heading = doc.add_paragraph("THỐNG KÊ", style='CustomHeading2')

        stats_table = doc.add_table(rows=2, cols=2)
        stats_table.style = 'Table Grid'

        # Header row
        hdr_cells = stats_table.rows[0].cells
        hdr_cells[0].text = "Loại"
        hdr_cells[1].text = "Số lượng"

        # Data rows
        stats_data = [
            ("Bằng chứng số", str(len(artifacts))),
            ("Kết quả phân tích", str(len(results)))
        ]

        for i, (label, value) in enumerate(stats_data, 1):
            row_cells = stats_table.rows[i].cells
            row_cells[0].text = label
            row_cells[1].text = value

        doc.add_page_break()

        # Conclusion Section
        heading = doc.add_paragraph("KẾT LUẬN CHÍNH", style='CustomHeading2')

        conclusion = doc.add_paragraph()
        conclusion.add_run(f"Vụ án đã được điều tra với {len(artifacts)} bằng chứng số và {len(results)} kết quả phân tích.").bold = True

        doc.add_paragraph("")
        doc.add_paragraph("Chuỗi bảo quản bằng chứng đã được duy trì theo đúng quy trình pháp y số.")

        return doc

    def _create_technical_docx(self, case_info, artifacts, results, activity_logs):
        """Create technical detailed Word document"""
        doc = Document()

        # Setup styles
        self._setup_word_styles(doc)

        # Header
        title = doc.add_paragraph("BÁO CÁO KỸ THUẬT", style='CustomHeading1')
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        case_title = doc.add_paragraph(f"Vụ án {case_info['case_id']}: {case_info['title']}")
        case_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph("")

        # Technical Details Section
        heading = doc.add_paragraph("CHI TIẾT KỸ THUẬT", style='CustomHeading2')

        details_table = doc.add_table(rows=4, cols=2)
        details_table.style = 'Table Grid'

        # Data rows
        rows_data = [
            ("Mã số vụ án", case_info['case_id']),
            ("Điều tra viên", case_info.get('full_name', 'Không xác định')),
            ("Ngày tạo", case_info.get('created_at', 'Không xác định')),
            ("Trạng thái", case_info.get('status', 'Không xác định'))
        ]

        for i, (label, value) in enumerate(rows_data):
            row_cells = details_table.rows[i].cells
            row_cells[0].text = label
            row_cells[1].text = value

        doc.add_page_break()

        # Evidence Collection Section
        heading = doc.add_paragraph("THU THẬP BẰNG CHỨNG", style='CustomHeading2')

        collection_table = doc.add_table(rows=3, cols=2)
        collection_table.style = 'Table Grid'

        # Data rows
        collection_data = [
            ("Tổng số bằng chứng", str(len(artifacts))),
            ("Tổng số kết quả", str(len(results))),
            ("Tổng số hoạt động", str(len(activity_logs)))
        ]

        for i, (label, value) in enumerate(collection_data):
            row_cells = collection_table.rows[i].cells
            row_cells[0].text = label
            row_cells[1].text = value

        doc.add_page_break()

        # Analysis Results Section
        heading = doc.add_paragraph("KẾT QUẢ PHÂN TÍCH", style='CustomHeading2')

        if results:
            results_table = doc.add_table(rows=1, cols=3)
            results_table.style = 'Table Grid'

            # Header row
            hdr_cells = results_table.rows[0].cells
            hdr_cells[0].text = "Mã số"
            hdr_cells[1].text = "Công cụ"
            hdr_cells[2].text = "Tóm tắt"

            # Data rows
            for result in results:
                row_cells = results_table.add_row().cells
                row_cells[0].text = str(result['result_id'])
                row_cells[1].text = result.get('tool_used', 'Không xác định')
                row_cells[2].text = result.get('summary', 'Không có thông tin')
        else:
            doc.add_paragraph("Chưa có kết quả phân tích nào.")

        return doc

    def _create_coc_docx(self, case_info, artifacts, activity_logs):
        """Create Chain of Custody Word document"""
        doc = Document()

        # Setup styles
        self._setup_word_styles(doc)

        # Header
        title = doc.add_paragraph("CHUỖI BẢO QUẢN BẰNG CHỨNG", style='CustomHeading1')
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        case_title = doc.add_paragraph(f"Vụ án {case_info['case_id']}: {case_info['title']}")
        case_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph("")

        # Evidence Chain Section
        heading = doc.add_paragraph("CHUỖI BẰNG CHỨNG", style='CustomHeading2')

        if artifacts:
            coc_table = doc.add_table(rows=1, cols=6)
            coc_table.style = 'Table Grid'

            # Header row
            hdr_cells = coc_table.rows[0].cells
            hdr_cells[0].text = "STT"
            hdr_cells[1].text = "Mô tả"
            hdr_cells[2].text = "Người thu thập"
            hdr_cells[3].text = "Ngày/Giờ"
            hdr_cells[4].text = "Mã băm SHA-256"
            hdr_cells[5].text = "Chữ ký"

            # Data rows
            for i, artifact in enumerate(artifacts, 1):
                row_cells = coc_table.add_row().cells
                row_cells[0].text = str(i)
                row_cells[1].text = artifact['name']
                row_cells[2].text = case_info.get('full_name', 'Không xác định')
                row_cells[3].text = artifact.get('collected_at', 'Không xác định')
                row_cells[4].text = artifact.get('sha256', 'Không có')
                row_cells[5].text = "_________________"
        else:
            doc.add_paragraph("Không tìm thấy bằng chứng nào.")

        doc.add_page_break()

        # Custody Transfer Log Section
        heading = doc.add_paragraph("NHẬT KÝ CHUYỂN GIAO BẢO QUẢN", style='CustomHeading2')

        transfer_table = doc.add_table(rows=1, cols=5)
        transfer_table.style = 'Table Grid'

        # Header row
        hdr_cells = transfer_table.rows[0].cells
        hdr_cells[0].text = "Từ"
        hdr_cells[1].text = "Đến"
        hdr_cells[2].text = "Ngày/Giờ"
        hdr_cells[3].text = "Mục đích"
        hdr_cells[4].text = "Chữ ký"

        # Data row
        row_cells = transfer_table.add_row().cells
        row_cells[0].text = case_info.get('full_name', 'Không xác định')
        row_cells[1].text = "Kho lưu trữ bằng chứng"
        row_cells[2].text = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        row_cells[3].text = "Lưu trữ an toàn"
        row_cells[4].text = "_________________"

        doc.add_page_break()

        # Integrity Verification Section
        heading = doc.add_paragraph("XÁC MINH TÍNH TOÀN VẸN", style='CustomHeading2')

        verification_items = [
            "Tất cả bằng chứng đã được xác minh bằng mã băm SHA-256.",
            "Chuỗi bảo quản đã được duy trì trong suốt quá trình điều tra.",
            f"Báo cáo được tạo: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}"
        ]

        for item in verification_items:
            p = doc.add_paragraph(item, style='List Bullet')

        return doc

    def _setup_word_styles(self, doc):
        """Setup Word document styles with error handling"""
        try:
            from docx.shared import Pt, RGBColor
            from docx.enum.style import WD_STYLE_TYPE

            styles = doc.styles

            # Check if styles already exist to avoid duplicates
            existing_styles = [s.name for s in styles]

            # Heading 1 style
            if 'CustomHeading1' not in existing_styles:
                h1_style = styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
                h1_style.font.size = Pt(18)
                h1_style.font.bold = True
                h1_style.font.color.rgb = RGBColor(31, 73, 125)

            # Heading 2 style
            if 'CustomHeading2' not in existing_styles:
                h2_style = styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
                h2_style.font.size = Pt(14)
                h2_style.font.bold = True
                h2_style.font.color.rgb = RGBColor(79, 129, 189)

            # Normal text style
            try:
                normal_style = styles['Normal']
                normal_style.font.size = Pt(11)
                normal_style.font.name = 'Times New Roman'
            except Exception as e:
                print(f"Warning: Could not set normal style: {e}")
                
        except Exception as e:
            print(f"Warning: Could not setup Word styles: {e}")
            # Continue without custom styles
    
    def _create_simplified_docx(self, case_info, artifacts, results, activity_logs, options=None):
        """Create simplified Word document that works reliably with selective sections"""
        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        
        # Default options if none provided
        if options is None:
            options = {
                "include_case_info": True,
                "include_evidence": True,
                "include_analysis": True,
                "include_activity": True,
                "include_coc": True,
                "include_summary": True
            }
        
        
        doc = Document()
        
        # Title (always included)
        title = doc.add_paragraph("BÁO CÁO ĐIỀU TRA SỐ TỔNG HỢP", style='Heading 1')
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        subtitle = doc.add_paragraph(f"Case ID: {case_info['case_id']} - {case_info['title']}")
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph("")
        
        # Case info section
        if options.get("include_case_info", True):
            doc.add_paragraph("THÔNG TIN VỤ ÁN", style='Heading 2')
            doc.add_paragraph(f"Mã số vụ án: {case_info['case_id']}")
            doc.add_paragraph(f"Tên vụ án: {case_info['title']}")
            doc.add_paragraph(f"Điều tra viên phụ trách: {case_info.get('full_name', 'Không xác định')}")
            doc.add_paragraph(f"Trạng thái hiện tại: {case_info.get('status', 'Không xác định')}")
            doc.add_paragraph(f"Ngày khởi tạo: {case_info.get('created_at', 'Không xác định')}")
            doc.add_paragraph(f"Thư mục lưu trữ: {case_info.get('archive_path', 'Không xác định')}")
            doc.add_paragraph("")
        
        # Artifacts section
        if options.get("include_evidence", True):
            doc.add_paragraph("BẰNG CHỨNG SỐ", style='Heading 2')
            if artifacts and len(artifacts) > 0:
                doc.add_paragraph(f"Tổng số bằng chứng thu thập được: {len(artifacts)} mục")
                for i, artifact in enumerate(artifacts, 1):
                    doc.add_paragraph(f"{i}. Tên bằng chứng: {artifact['name']}")
                    doc.add_paragraph(f"   - Loại bằng chứng: {artifact.get('evidence_type', 'Không xác định')}")
                    doc.add_paragraph(f"   - Kích thước tệp: {artifact.get('size', 0):,} bytes")
                    doc.add_paragraph(f"   - Thời gian thu thập: {artifact.get('collected_at', 'Không xác định')}")
                    doc.add_paragraph(f"   - Mã băm SHA-256: {artifact.get('sha256', 'Không có')[:32]}...")
            else:
                doc.add_paragraph("Chưa có bằng chứng số nào được thu thập trong vụ án này.")
            doc.add_paragraph("")
        
        # Results section
        if options.get("include_analysis", True):
            doc.add_paragraph("KẾT QUẢ PHÂN TÍCH", style='Heading 2')
            if results and len(results) > 0:
                doc.add_paragraph(f"Tổng số kết quả phân tích: {len(results)} kết quả")
                for i, result in enumerate(results, 1):
                    doc.add_paragraph(f"{i}. Công cụ sử dụng: {result.get('tool_used', 'Không xác định')}")
                    doc.add_paragraph(f"   - Thời gian thực hiện: {result.get('run_at', 'Không xác định')}")
                    doc.add_paragraph(f"   - Tóm tắt kết quả: {result.get('summary', 'Không có thông tin')}")
                    if result.get('result_path'):
                        doc.add_paragraph(f"   - Đường dẫn tệp kết quả: {result.get('result_path', 'Không có')}")
            else:
                doc.add_paragraph("Chưa có kết quả phân tích nào được thực hiện cho vụ án này.")
            doc.add_paragraph("")
        
        # Activity logs section
        if options.get("include_activity", True):
            doc.add_paragraph("NHẬT KÝ HOẠT ĐỘNG", style='Heading 2')
            if activity_logs and len(activity_logs) > 0:
                doc.add_paragraph(f"Tổng số hoạt động được ghi nhận: {len(activity_logs)} hoạt động")
                for i, log in enumerate(activity_logs, 1):
                    doc.add_paragraph(f"{i}. Hành động: {log.get('action', 'Không xác định')}")
                    doc.add_paragraph(f"   - Thời gian thực hiện: {log.get('timestamp', 'Không xác định')}")
                    doc.add_paragraph(f"   - Người thực hiện: {log.get('username', 'Không xác định')}")
                    if log.get('tool_used'):
                        doc.add_paragraph(f"   - Công cụ sử dụng: {log.get('tool_used', '')}")
                    if log.get('details'):
                        doc.add_paragraph(f"   - Chi tiết thêm: {log.get('details', '')}")
            else:
                doc.add_paragraph("Chưa có hoạt động nào được ghi nhận trong hệ thống.")
            doc.add_paragraph("")
        
        # Chain of custody section
        if options.get("include_coc", True):
            doc.add_paragraph("BIỂU MẪU CHUỖI BẢO QUẢN BẰNG CHỨNG", style='Heading 2')
            doc.add_paragraph("")
            
            # Evidence/Property Custody Document Header
            doc.add_paragraph("BIÊN BẢN BẢO QUẢN BẰNG CHỨNG VẬT CHỨNG", style='Heading 3')
            
            # Case information table
            case_info_table = doc.add_table(rows=3, cols=4)
            case_info_table.style = 'Table Grid'
            
            # Row 1: Tracking number and Case ID
            row1_cells = case_info_table.rows[0].cells
            row1_cells[0].text = "SỐ THEO DÕI"
            row1_cells[1].text = f"VỤ ÁN-{case_info['case_id']}-{datetime.now().strftime('%Y%m%d')}"
            row1_cells[2].text = "MÃ SỐ VỤ ÁN"
            row1_cells[3].text = str(case_info['case_id'])
            
            # Row 2: Receiving activity and Location
            row2_cells = case_info_table.rows[1].cells
            row2_cells[0].text = "HOẠT ĐỘNG TIẾP NHẬN"
            row2_cells[1].text = "Điều tra pháp y số"
            row2_cells[2].text = "ĐỊA ĐIỂM"
            row2_cells[3].text = "Phòng lab pháp y số"
            
            # Row 3: Name/address and other info
            row3_cells = case_info_table.rows[2].cells
            row3_cells[0].text = "TÊN, VĂN PHÒNG VÀ CHỨC DANH NGƯỜI GIAO BẰNG CHỨNG"
            row3_cells[1].text = f"Chủ sở hữu: {case_info.get('full_name', 'Không xác định')}"
            row3_cells[2].text = "ĐỊA CHỈ (Bao gồm mã bưu điện)"
            row3_cells[3].text = case_info.get('archive_path', 'Không xác định')
            
            doc.add_paragraph("")
            
            # Location and reason table
            location_table = doc.add_table(rows=1, cols=3)
            location_table.style = 'Table Grid'
            
            loc_cells = location_table.rows[0].cells
            loc_cells[0].text = "ĐỊA ĐIỂM THU THẬP"
            loc_cells[1].text = "LÝ DO THU THẬP"
            loc_cells[2].text = "NGÀY/GIỜ THU THẬP"
            
            # Add data row
            data_row = location_table.add_row()
            data_cells = data_row.cells
            data_cells[0].text = case_info.get('archive_path', 'Không xác định')
            data_cells[1].text = "Điều tra pháp y số"
            data_cells[2].text = case_info.get('created_at', 'Không xác định')
            
            doc.add_paragraph("")
            
            # Evidence items table
            doc.add_paragraph("MÔ TẢ VẬT CHỨNG", style='Heading 3')
            
            evidence_table = doc.add_table(rows=1, cols=3)
            evidence_table.style = 'Table Grid'
            
            # Header
            ev_hdr = evidence_table.rows[0].cells
            ev_hdr[0].text = "STT"
            ev_hdr[1].text = "SỐ LƯỢNG"
            ev_hdr[2].text = "MÔ TẢ VẬT CHỨNG\n(Bao gồm model, số serial, tình trạng và các dấu hiệu đặc biệt)"
            
            # Add evidence items
            if artifacts and len(artifacts) > 0:
                for i, artifact in enumerate(artifacts, 1):
                    ev_row = evidence_table.add_row()
                    ev_cells = ev_row.cells
                    ev_cells[0].text = str(i)
                    ev_cells[1].text = "1"
                    
                    description = f"Bằng chứng số: {artifact['name']}\n"
                    description += f"Loại: {artifact.get('evidence_type', 'Không xác định')}\n"
                    description += f"Kích thước: {artifact.get('size', 0):,} bytes\n"
                    description += f"Mã băm SHA256: {artifact.get('sha256', 'Không có')[:32]}...\n" if artifact.get('sha256', 'Không có') != 'Không có' else "Mã băm SHA256: Không có\n"
                    description += f"Đường dẫn: {artifact.get('source_path', 'Không xác định')}"
                    
                    ev_cells[2].text = description
            else:
                ev_row = evidence_table.add_row()
                ev_cells = ev_row.cells
                ev_cells[0].text = "1"
                ev_cells[1].text = "0"
                ev_cells[2].text = "Chưa có bằng chứng số nào được xử lý"
            
            doc.add_paragraph("")
            
            # Chain of Custody tracking table
            doc.add_paragraph("CHUỖI BẢO QUẢN BẰNG CHỨNG", style='Heading 3')
            
            coc_table = doc.add_table(rows=1, cols=6)
            coc_table.style = 'Table Grid'
            
            # Header
            coc_hdr = coc_table.rows[0].cells
            coc_hdr[0].text = "STT VẬT CHỨNG"
            coc_hdr[1].text = "NGÀY THÁNG"
            coc_hdr[2].text = "NGƯỜI GIAO\nCHỮ KÝ"
            coc_hdr[3].text = "NGƯỜI NHẬN\nCHỮ KÝ"
            coc_hdr[4].text = "MỤC ĐÍCH CHUYỂN GIAO"
            coc_hdr[5].text = "TÊN, CẤP BẬC HOẶC CHỨC DANH"
            
            # Add initial custody entry
            coc_row1 = coc_table.add_row()
            coc_cells1 = coc_row1.cells
            coc_cells1[0].text = "TẤT CẢ"
            coc_cells1[1].text = case_info.get('created_at', datetime.now().strftime('%Y-%m-%d'))
            coc_cells1[2].text = "THU THẬP BAN ĐẦU"
            coc_cells1[3].text = case_info.get('full_name', 'Chuyên viên pháp y số')
            coc_cells1[4].text = "Thu thập bằng chứng"
            coc_cells1[5].text = case_info.get('full_name', 'Không xác định')
            
            # Add analysis entry
            coc_row2 = coc_table.add_row()
            coc_cells2 = coc_row2.cells
            coc_cells2[0].text = "TẤT CẢ"
            coc_cells2[1].text = datetime.now().strftime('%Y-%m-%d')
            coc_cells2[2].text = case_info.get('full_name', 'Không xác định')
            coc_cells2[3].text = "Hệ thống pháp y số"
            coc_cells2[4].text = "Phân tích pháp y số"
            coc_cells2[5].text = "Phân tích toàn diện"
            
            # Add empty rows for future custody transfers
            for i in range(3):
                coc_row = coc_table.add_row()
                coc_cells = coc_row.cells
                for cell in coc_cells:
                    cell.text = ""
            
            doc.add_paragraph("")
            
            # Final disposal section
            doc.add_paragraph("HÀNH ĐỘNG XỬ LÝ CUỐI CÙNG", style='Heading 3')
            
            disposal_table = doc.add_table(rows=4, cols=1)
            disposal_table.style = 'Table Grid'
            
            disposal_table.rows[0].cells[0].text = f"TRẢ LẠI CHO CHỦ SỞ HỮU HOẶC KHÁC (Tên/Tổ chức): {case_info.get('full_name', 'Không xác định')}"
            disposal_table.rows[1].cells[0].text = "HỦY BỎ: _______________"
            disposal_table.rows[2].cells[0].text = "KHÁC (Ghi rõ): Phân tích số hoàn tất - Bằng chứng được trả lại"
            disposal_table.rows[3].cells[0].text = "CƠ QUAN THẨM QUYỀN XỬ LÝ CUỐI: Phòng Pháp y Số"
            
            doc.add_paragraph("")
            
            # Certification
            cert_text = f"TÔI XÁC NHẬN RẰNG THÔNG TIN TRÊN LÀ BẢN GHI CHÍNH XÁC VỀ VIỆC BẢO QUẢN TRONG QUÁ TRÌNH ĐIỀU TRA/KHẢO SÁT CÁC VẬT CHỨNG: Bằng chứng số\n"
            cert_text += f"VÀ CÁC VẬT CHỨNG ĐÃ ĐƯỢC BẢO QUẢN VÀ XỬ LÝ ĐÚNG QUY ĐỊNH NHƯ ĐƯỢC CHỈ RA Ở TRÊN.\n\n"
            cert_text += f"CẦN THIẾT LÀM BẰNG CHỨNG VÀ CÓ THỂ ĐƯỢC XỬ LÝ NHƯ CHỈ RA Ở TRÊN. (Nếu vật phẩm phải được giữ lại thì không ký mà giải thích trong văn bản riêng.)\n\n"
            cert_text += f"Điều tra viên: {case_info.get('full_name', 'Không xác định')}\n"
            cert_text += f"Chữ ký: ___________________________ Ngày: {datetime.now().strftime('%Y-%m-%d')}\n\n"
            cert_text += f"NHÂN CHỨNG VIỆC HỦY BỎ BẰNG CHỨNG\n"
            cert_text += f"CÁC VẬT PHẨM ĐƯỢC LIỆT KÊ TẠI MỤC SỐ: Tệp bằng chứng số\n"
            cert_text += f"ĐÃ ĐƯỢC NGƯỜI BẢO QUẢN BẰNG CHỨNG HỦY BỎ TRƯỚC MẶT TÔI, VÀO NGÀY NHƯ CHỈ RA Ở TRÊN.\n\n"
            cert_text += f"Nhân chứng: ___________________________\n"
            cert_text += f"Chữ ký: ___________________________ Ngày: ___________"
            
            doc.add_paragraph(cert_text)
            doc.add_paragraph("")
        
        # Summary and recommendations section
        if options.get("include_summary", True):
            doc.add_paragraph("TÓM TẮT VÀ KHUYẾN NGHỊ", style='Heading 2')
            
            # Statistics
            doc.add_paragraph("Thống kê tổng quan:")
            doc.add_paragraph(f"- Tổng số bằng chứng số: {len(artifacts) if artifacts else 0} mục")
            doc.add_paragraph(f"- Tổng số kết quả phân tích: {len(results) if results else 0} kết quả")
            doc.add_paragraph(f"- Tổng số hoạt động ghi nhận: {len(activity_logs) if activity_logs else 0} hoạt động")
            doc.add_paragraph("")
            
            # Recommendations
            doc.add_paragraph("Khuyến nghị:")
            if len(artifacts) == 0:
                doc.add_paragraph("- Cần thu thập thêm bằng chứng số để hỗ trợ quá trình điều tra")
            if len(results) == 0:
                doc.add_paragraph("- Cần thực hiện phân tích với các công cụ pháp y số chuyên dụng")
            
            doc.add_paragraph("- Hoàn thiện và duy trì chuỗi bảo quản bằng chứng theo quy định")
            doc.add_paragraph("- Chuẩn bị báo cáo chi tiết để trình cơ quan có thẩm quyền")
            doc.add_paragraph("- Lưu trữ an toàn và bảo mật tất cả bằng chứng số thu thập được")
            doc.add_paragraph("")
        
        # Footer (always included)
        doc.add_paragraph("─" * 60)
        footer_text = f"Báo cáo được tạo vào: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}"
        footer_text += f"\nHệ thống pháp y số - Windows Forensic System"
        footer_text += f"\nĐiều tra viên phụ trách: {case_info.get('full_name', 'Không xác định')}"
        doc.add_paragraph(footer_text)
        
        return doc

class Report(QWidget):
    def __init__(self, main_window=None):
        super(Report, self).__init__()
        self.main_window = main_window
        self.ui = Ui_Form()
        self.ui.setupUi(self)
        self.setup_ui()
        self.setup_connections()
        
        self.db = DatabaseManager()
        self.current_case_id = None
        self.report_generator = None
        
    def setup_ui(self):
        """Setup additional UI components and styling"""
        # Set table headers
        self.ui.reportsTable.setColumnCount(5)
        self.ui.reportsTable.setHorizontalHeaderLabels([
            "ID", "Loại", "Định dạng", "Ngày tạo", "Hash SHA-256"
        ])
        
        # Set table properties
        header = self.ui.reportsTable.horizontalHeader()
        header.setSectionResizeMode(0, QHeaderView.ResizeToContents)
        header.setSectionResizeMode(1, QHeaderView.Stretch)
        header.setSectionResizeMode(2, QHeaderView.ResizeToContents)
        header.setSectionResizeMode(3, QHeaderView.ResizeToContents)
        header.setSectionResizeMode(4, QHeaderView.ResizeToContents)
        
        # Set initial state
        self.ui.generateButton.setEnabled(False)
        
    def setup_connections(self):
        """Setup signal connections"""
        self.ui.generateButton.clicked.connect(self.generate_report)
        self.ui.refreshButton.clicked.connect(self.refresh_cases)
        self.ui.caseCombo.currentTextChanged.connect(self.on_case_changed)
        
        # Add context menu for reports table
        self.ui.reportsTable.setContextMenuPolicy(Qt.CustomContextMenu)
        self.ui.reportsTable.customContextMenuRequested.connect(self.show_reports_context_menu)
        
    def refresh_cases(self):
        """Refresh case list from database"""
        try:
            if not self.db.connect():
                QMessageBox.warning(self, "Lỗi", "Không thể kết nối database!")
                return
                
            # Sửa method name từ get_all_cases thành get_cases
            cases = self.db.get_cases()
            self.ui.caseCombo.clear()
            self.ui.caseCombo.addItem("-- Chọn vụ án --", None)
            
            for case in cases:
                display_text = f"{case['case_id']} - {case['title']}"
                self.ui.caseCombo.addItem(display_text, case['case_id'])
                
            # Nếu có main_window, lấy case đang được chọn
            if self.main_window and hasattr(self.main_window, 'current_case_id'):
                self.select_current_working_case()
                
        except Exception as e:
            print(f"Error refreshing cases: {e}")
            QMessageBox.warning(self, "Lỗi", f"Không thể tải danh sách vụ án: {str(e)}")
        finally:
            self.db.disconnect()
            
    def select_current_working_case(self):
        """Tự động chọn case đang được làm việc từ main window"""
        if self.main_window and hasattr(self.main_window, 'current_case_id'):
            current_case_id = self.main_window.current_case_id
            if current_case_id:
                # Tìm và chọn case trong combo box
                for i in range(self.ui.caseCombo.count()):
                    if self.ui.caseCombo.itemData(i) == current_case_id:
                        self.ui.caseCombo.setCurrentIndex(i)
                        self.current_case_id = current_case_id
                        self.ui.generateButton.setEnabled(True)
                        self.load_case_reports(current_case_id)
                        print(f"Auto-selected current working case: {current_case_id}")
                        break
                        
    def on_case_changed(self):
        """Handle case selection change"""
        case_id = self.ui.caseCombo.currentData()
        self.current_case_id = case_id
        
        if case_id:
            self.load_case_reports(case_id)
            self.ui.generateButton.setEnabled(True)
            
            # Cập nhật current_case_id trong main window nếu có
            if self.main_window and hasattr(self.main_window, 'current_case_id'):
                self.main_window.current_case_id = case_id
                print(f"Updated main window current_case_id to: {case_id}")
        else:
            self.ui.generateButton.setEnabled(False)
            self.ui.reportsTable.setRowCount(0)
            
    def validate_case_data(self):
        """Kiểm tra xem case có dữ liệu để tạo báo cáo không"""
        try:
            if not self.db.connect():
                QMessageBox.warning(self, "Lỗi", "Không thể kết nối database!")
                return False
            
            # Lấy thông tin case
            case_info = self.db.get_case_with_investigator(self.current_case_id)
            if not case_info:
                QMessageBox.warning(self, "Lỗi", f"Không tìm thấy thông tin case ID: {self.current_case_id}")
                return False
            
            # Lấy số lượng artifacts và results
            artifacts = self.db.get_artifacts_by_case(self.current_case_id)
            results = self.db.get_results_by_case(self.current_case_id)
            
            artifact_count = len(artifacts) if artifacts else 0
            result_count = len(results) if results else 0
            
            # Hiển thị thông tin case
            info_msg = (
                f"📋 Thông tin Case:\n\n"
                f"🆔 ID: {case_info['case_id']}\n"
                f"📝 Tên: {case_info['title']}\n"
                f"👨‍💼 Điều tra viên: {case_info.get('full_name', 'N/A')}\n"
                f"📅 Ngày tạo: {case_info.get('created_at', 'N/A')}\n"
                f"📊 Trạng thái: {case_info.get('status', 'N/A')}\n\n"
                f"📁 Bằng chứng số: {artifact_count}\n"
                f"🔬 Kết quả phân tích: {result_count}\n\n"
            )
            
            if artifact_count == 0 and result_count == 0:
                info_msg += (
                    "⚠️ Case này chưa có dữ liệu!\n\n"
                    "Để tạo báo cáo có ý nghĩa, bạn cần:\n"
                    "1. Thu thập bằng chứng số (artifacts)\n"
                    "2. Chạy phân tích với các công cụ forensics\n"
                    "3. Sau đó mới xuất báo cáo\n\n"
                    "Bạn có muốn tiếp tục tạo báo cáo trống không?"
                )
                
                reply = QMessageBox.question(
                    self, 
                    "⚠️ Case chưa có dữ liệu", 
                    info_msg,
                    QMessageBox.Yes | QMessageBox.No,
                    QMessageBox.No
                )
                return reply == QMessageBox.Yes
            else:
                info_msg += "✅ Case có đủ dữ liệu để tạo báo cáo!"
                
                reply = QMessageBox.information(
                    self,
                    "📋 Xác nhận tạo báo cáo",
                    info_msg + "\n\nTiếp tục tạo báo cáo?",
                    QMessageBox.Yes | QMessageBox.No,
                    QMessageBox.Yes
                )
                return reply == QMessageBox.Yes
                
        except Exception as e:
            print(f"Error validating case data: {e}")
            QMessageBox.warning(self, "Lỗi", f"Không thể kiểm tra dữ liệu case: {str(e)}")
            return False
        finally:
            self.db.disconnect()

    def load_case_reports(self, case_id):
        """Load reports for selected case"""
        try:
            if not self.db.connect():
                return
                
            reports = self.db.get_reports_by_case(case_id)
            
            # Clear table first
            self.ui.reportsTable.setRowCount(0)
            self.ui.reportsTable.clearContents()
            
            if not reports:
                print(f"No reports found for case {case_id}")
                return
                
            self.ui.reportsTable.setRowCount(len(reports))
            
            for row, report in enumerate(reports):
                # Store report_id as item data for later use
                id_item = QTableWidgetItem(str(report['report_id']))
                id_item.setData(Qt.UserRole, report['report_id'])
                self.ui.reportsTable.setItem(row, 0, id_item)
                
                self.ui.reportsTable.setItem(row, 1, QTableWidgetItem('Báo cáo tổng hợp'))
                self.ui.reportsTable.setItem(row, 2, QTableWidgetItem('DOCX'))
                self.ui.reportsTable.setItem(row, 3, QTableWidgetItem(report.get('created_at', 'N/A')))
                
                # Truncate hash for display
                hash_value = report.get('sha256', 'N/A')
                if hash_value != 'N/A' and len(hash_value) > 16:
                    hash_value = hash_value[:16] + '...'
                self.ui.reportsTable.setItem(row, 4, QTableWidgetItem(hash_value))
                
        except Exception as e:
            print(f"Error loading reports: {e}")
        finally:
            self.db.disconnect()
            
    def show_reports_context_menu(self, position):
        """Show context menu for reports table"""
        if self.ui.reportsTable.itemAt(position) is None:
            return
            
        from PyQt5.QtWidgets import QMenu, QAction
        
        menu = QMenu(self)
        
        # Add actions
        refresh_action = QAction("🔄 Làm mới", self)
        refresh_action.triggered.connect(self.refresh_reports_table)
        menu.addAction(refresh_action)
        
        delete_action = QAction("🗑️ Xóa báo cáo", self)
        delete_action.triggered.connect(self.delete_selected_report)
        menu.addAction(delete_action)
        
        menu.addSeparator()
        
        open_action = QAction("📂 Mở file", self)
        open_action.triggered.connect(self.open_selected_report)
        menu.addAction(open_action)
        
        # Show menu
        menu.exec_(self.ui.reportsTable.mapToGlobal(position))
        
    def refresh_reports_table(self):
        """Refresh reports table for current case"""
        if self.current_case_id:
            print(f"Refreshing reports table for case {self.current_case_id}")
            self.load_case_reports(self.current_case_id)
        else:
            self.ui.reportsTable.setRowCount(0)
            self.ui.reportsTable.clearContents()
            
    def delete_selected_report(self):
        """Delete selected report"""
        current_row = self.ui.reportsTable.currentRow()
        if current_row < 0:
            QMessageBox.warning(self, "Cảnh báo", "Vui lòng chọn báo cáo cần xóa!")
            return
            
        # Get report ID
        id_item = self.ui.reportsTable.item(current_row, 0)
        if not id_item:
            return
            
        report_id = id_item.data(Qt.UserRole)
        report_name = id_item.text()
        
        # Confirm deletion
        reply = QMessageBox.question(
            self,
            "❓ Xác nhận xóa",
            f"Bạn có chắc chắn muốn xóa báo cáo ID {report_name}?\n\nHành động này không thể hoàn tác!",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No
        )
        
        if reply == QMessageBox.Yes:
            try:
                if not self.db.connect():
                    QMessageBox.warning(self, "Lỗi", "Không thể kết nối database!")
                    return
                    
                # Delete from database
                cursor = self.db.connection.cursor()
                cursor.execute("DELETE FROM reports WHERE report_id = ?", (report_id,))
                self.db.connection.commit()
                
                QMessageBox.information(self, "✅ Thành công", f"Đã xóa báo cáo ID {report_name}")
                
                # Refresh table
                self.refresh_reports_table()
                
            except Exception as e:
                print(f"Error deleting report: {e}")
                QMessageBox.critical(self, "❌ Lỗi", f"Không thể xóa báo cáo: {str(e)}")
            finally:
                self.db.disconnect()
                
    def open_selected_report(self):
        """Open selected report file"""
        current_row = self.ui.reportsTable.currentRow()
        if current_row < 0:
            QMessageBox.warning(self, "Cảnh báo", "Vui lòng chọn báo cáo cần mở!")
            return
            
        # Get report ID and find file path
        id_item = self.ui.reportsTable.item(current_row, 0)
        if not id_item:
            return
            
        report_id = id_item.data(Qt.UserRole)
        
        try:
            if not self.db.connect():
                QMessageBox.warning(self, "Lỗi", "Không thể kết nối database!")
                return
                
            # Get file path from database
            cursor = self.db.connection.cursor()
            cursor.execute("SELECT file_path FROM reports WHERE report_id = ?", (report_id,))
            result = cursor.fetchone()
            
            if result and result[0]:
                file_path = result[0]
                if os.path.exists(file_path):
                    try:
                        os.startfile(file_path)  # Windows
                    except:
                        QMessageBox.warning(self, "Lỗi", f"Không thể mở file: {file_path}")
                else:
                    QMessageBox.warning(self, "Lỗi", f"File không tồn tại: {file_path}")
            else:
                QMessageBox.warning(self, "Lỗi", "Không tìm thấy đường dẫn file!")
                
        except Exception as e:
            print(f"Error opening report: {e}")
            QMessageBox.critical(self, "❌ Lỗi", f"Không thể mở báo cáo: {str(e)}")
        finally:
            self.db.disconnect()
            
    def generate_report(self):
        """Generate comprehensive report with selected sections"""
        if not self.current_case_id:
            QMessageBox.warning(self, "Cảnh báo", "Vui lòng chọn vụ án trước!")
            return
            
        # Kiểm tra xem case có dữ liệu không
        if not self.validate_case_data():
            return
            
        # Fixed report type - always comprehensive
        report_type = "comprehensive"
        
        # Get section options from checkboxes
        options = {
            "include_case_info": self.ui.caseInfoCheckbox.isChecked(),
            "include_evidence": self.ui.evidenceCheckbox.isChecked(),
            "include_analysis": self.ui.analysisCheckbox.isChecked(),
            "include_activity": self.ui.activityCheckbox.isChecked(),
            "include_coc": self.ui.cocCheckbox.isChecked(),
            "include_summary": self.ui.summaryCheckbox.isChecked()
        }
        
        # Start report generation
        self.ui.generateButton.setEnabled(False)
        self.ui.progressBar.setVisible(True)
        self.ui.progressBar.setValue(0)
        
        # Create and start report generator thread
        self.report_generator = ReportGenerator(self.current_case_id, report_type, options)
        self.report_generator.progress_updated.connect(self.update_progress)
        self.report_generator.report_generated.connect(self.on_report_generated)
        self.report_generator.start()
        
        print(f"🚀 Bắt đầu tạo báo cáo {report_type} cho Case ID: {self.current_case_id}")
        
        # Simulate progress updates
        self.progress_timer = QTimer()
        self.progress_timer.timeout.connect(self.simulate_progress)
        self.progress_timer.start(100)
        
    def simulate_progress(self):
        """Simulate progress updates"""
        current = self.ui.progressBar.value()
        if current < 90:
            self.ui.progressBar.setValue(current + 5)
            
    def update_progress(self, value, message):
        """Update progress bar"""
        self.ui.progressBar.setValue(value)
        
    def on_report_generated(self, file_path, report_type):
        """Handle report generation completion"""
        self.progress_timer.stop()
        self.ui.progressBar.setVisible(False)
        self.ui.generateButton.setEnabled(True)
        
        if file_path and report_type != "error":
            # Lấy thông tin file
            import os
            file_size = os.path.getsize(file_path) if os.path.exists(file_path) else 0
            file_name = os.path.basename(file_path)
            folder_path = os.path.dirname(file_path)
            
            success_msg = (
                f"✅ Báo cáo Word đã được tạo thành công!\n\n"
                f"📄 Tên file: {file_name}\n"
                f"📊 Loại báo cáo: {report_type}\n"
                f"📏 Kích thước: {file_size:,} bytes\n"
                f"📂 Thư mục: {folder_path}\n\n"
                f"🔐 Báo cáo đã được lưu vào database với hash SHA-256 để đảm bảo tính toàn vẹn.\n\n"
                f"💡 Bạn có thể:\n"
                f"• Mở file Word để xem nội dung\n"
                f"• Tạo thêm các loại báo cáo khác\n"
                f"• Xem lịch sử báo cáo trong bảng bên dưới"
            )
            
            QMessageBox.information(self, "🎉 Tạo báo cáo thành công", success_msg)
            
            # Refresh reports table
            if self.current_case_id:
                self.load_case_reports(self.current_case_id)
                
            # Hỏi xem có muốn mở file không
            reply = QMessageBox.question(
                self,
                "📂 Mở file báo cáo",
                f"Bạn có muốn mở file báo cáo vừa tạo không?\n\n{file_name}",
                QMessageBox.Yes | QMessageBox.No,
                QMessageBox.Yes
            )
            
            if reply == QMessageBox.Yes:
                try:
                    os.startfile(file_path)  # Windows
                except:
                    try:
                        import subprocess
                        subprocess.run(['xdg-open', file_path])  # Linux
                    except:
                        try:
                            subprocess.run(['open', file_path])  # macOS
                        except:
                            print(f"Không thể mở file: {file_path}")
        else:
            error_msg = (
                "❌ Không thể tạo báo cáo Word!\n\n"
                "Có thể do các nguyên nhân sau:\n"
                "• Thiếu thư viện python-docx\n"
                "• Không có quyền ghi file\n"
                "• Lỗi database\n"
                "• Dữ liệu case không hợp lệ\n\n"
                "💡 Hướng dẫn khắc phục:\n"
                "1. Cài đặt: pip install python-docx\n"
                "2. Kiểm tra quyền ghi thư mục\n"
                "3. Đảm bảo case có dữ liệu\n"
                "4. Kiểm tra log console để biết chi tiết"
            )
            
            QMessageBox.critical(self, "❌ Lỗi tạo báo cáo", error_msg)
            
    def set_case_id(self, case_id):
        """Set current case ID from external source"""
        if case_id:
            # Find and select the case in combo box
            for i in range(self.ui.caseCombo.count()):
                if self.ui.caseCombo.itemData(i) == case_id:
                    self.ui.caseCombo.setCurrentIndex(i)
                    break
                    
    def showEvent(self, event):
        """Handle show event - refresh cases when tab becomes visible"""
        super().showEvent(event)
        self.refresh_cases()