# -*- coding: utf-8 -*-

import os
import sys
import csv
import json
import subprocess
import hashlib
import traceback
from datetime import datetime, timedelta
from pathlib import Path
from collections import defaultdict

# Phân tích Registry (Windows Registry)
try:
    import Registry.Registry as Registry
    REGISTRY_AVAILABLE = True
except ImportError:
    Registry = None
    REGISTRY_AVAILABLE = False
    print("Cảnh báo: Chưa cài đặt thư viện python-registry. Cài đặt bằng: pip install python-registry")

# PyQt5 - Framework GUI
from PyQt5.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QMessageBox, QFileDialog,
    QTableWidgetItem, QProgressDialog, QApplication,
    QComboBox, QLabel, QPushButton, QTextEdit, QTabWidget,
    QAbstractItemView, QMenu, QAction, QListWidgetItem, QTreeWidgetItem,
    QHeaderView, QDialog, QCheckBox, QScrollArea, QFrame, QSizePolicy
)
from PyQt5.QtCore import Qt, QThread, pyqtSignal, QTimer, QDateTime, QModelIndex
from PyQt5.QtGui import QStandardItemModel, QStandardItem, QIcon, QColor, QFont

# Import giao diện người dùng - LƯU Ý: import đúng tên file UI
from views.pages.analysis_ui.registry_analysis_ui import Ui_RegistryAnalysisWidget

# Word document generation (python-docx)
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    Document = None
    WD_PARAGRAPH_ALIGNMENT = None
    DOCX_AVAILABLE = False
    print("Cảnh báo: Chưa cài đặt thư viện python-docx. Cài đặt bằng: pip install python-docx")

# ============= Các Hàm Tiện Ích (giữ nguyên từ code cũ) =============

def format_as_hex(data):
    """Định dạng dữ liệu thành hex dump với preview ASCII."""
    if not data:
        return "Không có dữ liệu"

    try:
        if isinstance(data, str):
            data = data.encode("utf-8", errors="ignore")

        lines = []
        for offset in range(0, len(data), 16):
            chunk = data[offset:offset+16]
            hex_bytes = ' '.join(f'{b:02x}' for b in chunk)
            ascii_part = ''.join(chr(b) if 32 <= b <= 126 else '.' for b in chunk)
            lines.append(f"{offset:08X}  {hex_bytes:<48} |{ascii_part}|")
        return "\n".join(lines)
    except Exception as e:
        return f"Lỗi định dạng hex: {str(e)}"

def decode_registry_data(data, format_type):
    """Giải mã dữ liệu registry dựa trên định dạng được chọn."""
    if not data:
        return "Không có dữ liệu để giải mã"

    try:
        if format_type == "Auto-detect":
            return auto_decode_data(data)
        elif format_type == "UTF-16 String":
            if isinstance(data, bytes):
                return data.decode("utf-16le", errors="ignore").rstrip('\x00')
            return str(data)
        elif format_type == "UTF-8 String":
            if isinstance(data, bytes):
                return data.decode("utf-8", errors="ignore").rstrip('\x00')
            return str(data)
        elif format_type == "DWORD (32-bit)":
            return decode_dword(data)
        elif format_type == "QWORD (64-bit)":
            return decode_qword(data)
        elif format_type == "Windows FILETIME":
            return decode_filetime(data)
        elif format_type == "SID":
            return decode_sid(data)
        elif format_type == "GUID":
            return decode_guid(data)
        else:
            return str(data)
    except Exception as e:
        return f"Lỗi giải mã: {str(e)}"

def auto_decode_data(data):
    """Tự động phát hiện định dạng và giải mã dữ liệu."""
    if not data:
        return "Không có dữ liệu"

    try:
        # Thử dưới dạng số
        if isinstance(data, int):
            return f"Số nguyên: {data}\nHex: 0x{data:X}"

        # Thử dưới dạng chuỗi
        if isinstance(data, str):
            return f"Chuỗi: {data}"

        # Thử dưới dạng bytes
        if isinstance(data, bytes):
            # Thử UTF-16 trước (thường dùng trong Windows)
            try:
                utf16 = data.decode("utf-16le", errors="strict").rstrip('\x00')
                if utf16.isprintable() or '\n' in utf16:
                    return f"Chuỗi UTF-16: {utf16}"
            except:
                pass

            # Thử UTF-8
            try:
                utf8 = data.decode("utf-8", errors="strict").rstrip('\x00')
                if utf8.isprintable():
                    return f"Chuỗi UTF-8: {utf8}"
            except:
                pass

            # Kiểm tra các pattern thường gặp
            if len(data) == 4:
                value = int.from_bytes(data, byteorder="little")
                return f"DWORD: {value}\nHex: 0x{value:08X}"
            elif len(data) == 8:
                value = int.from_bytes(data, byteorder="little")
                return f"QWORD: {value}\nHex: 0x{value:016X}"

            # Mặc định hiển thị hex
            return format_as_hex(data)

        return str(data)
    except Exception as e:
        return f"Lỗi tự động giải mã: {str(e)}"

def decode_dword(data):
    """Giải mã giá trị DWORD (32-bit)."""
    try:
        if isinstance(data, int):
            return f"DWORD: {data}\nHex: 0x{data:08X}\nNhị phân: {bin(data)}"
        elif isinstance(data, bytes) and len(data) == 4:
            value = int.from_bytes(data, byteorder="little")
            return f"DWORD: {value}\nHex: 0x{value:08X}\nNhị phân: {bin(value)}"
        return str(data)
    except Exception as e:
        return f"Lỗi giải mã DWORD: {str(e)}"

def decode_qword(data):
    """Giải mã giá trị QWORD (64-bit)."""
    try:
        if isinstance(data, int):
            return f"QWORD: {data}\nHex: 0x{data:016X}"
        elif isinstance(data, bytes) and len(data) == 8:
            value = int.from_bytes(data, byteorder="little")
            return f"QWORD: {value}\nHex: 0x{value:016X}"
        return str(data)
    except Exception as e:
        return f"Lỗi giải mã QWORD: {str(e)}"

def decode_filetime(data):
    """Giải mã Windows FILETIME thành datetime có thể đọc được."""
    try:
        if isinstance(data, bytes) and len(data) == 8:
            filetime = int.from_bytes(data, byteorder="little")
            if filetime == 0:
                return "FILETIME: Chưa thiết lập (0)"
            # Chuyển từ Windows epoch (1601) sang Unix epoch
            dt = datetime(1601, 1, 1) + timedelta(microseconds=filetime/10)
            return f"FILETIME: {dt.strftime('%Y-%m-%d %H:%M:%S.%f')[:-3]}\nThô: {filetime}"
        return f"Dữ liệu FILETIME không hợp lệ"
    except Exception as e:
        return f"Lỗi giải mã FILETIME: {str(e)}"

def decode_sid(data):
    """Giải mã Windows Security Identifier."""
    try:
        if isinstance(data, bytes) and len(data) >= 8:
            revision = data[0]
            sub_auth_count = data[1]
            authority = int.from_bytes(data[2:8], byteorder="big")

            sid_string = f"S-{revision}-{authority}"

            for i in range(sub_auth_count):
                if 8 + (i * 4) + 4 <= len(data):
                    sub_auth = int.from_bytes(data[8+(i*4):8+(i*4)+4], byteorder="little")
                    sid_string += f"-{sub_auth}"

            return f"SID: {sid_string}"
        return "Dữ liệu SID không hợp lệ"
    except Exception as e:
        return f"Lỗi giải mã SID: {str(e)}"

def decode_guid(data):
    """Giải mã GUID/UUID."""
    try:
        if isinstance(data, bytes) and len(data) == 16:
            # Định dạng: {XXXXXXXX-XXXX-XXXX-XXXX-XXXXXXXXXXXX}
            p1 = data[0:4][::-1].hex()
            p2 = data[4:6][::-1].hex()
            p3 = data[6:8][::-1].hex()
            p4 = data[8:10].hex()
            p5 = data[10:16].hex()

            guid = f"{{{p1}-{p2}-{p3}-{p4}-{p5}}}".upper()
            return f"GUID: {guid}"
        return "Dữ liệu GUID không hợp lệ"
    except Exception as e:
        return f"Lỗi giải mã GUID: {str(e)}"

# ============= Thread Phân Tích Registry (giữ nguyên) =============

class RegistryAnalysisThread(QThread):
    """Thread để chạy phân tích RECmd mà không chặn UI."""

    progress_updated = pyqtSignal(int)
    status_updated = pyqtSignal(str)
    analysis_completed = pyqtSignal(dict)
    error_occurred = pyqtSignal(str)

    def __init__(self, recmd_path, batch_file, hive_files, output_dir):
        super().__init__()
        self.recmd_path = recmd_path
        self.batch_file = batch_file
        self.hive_files = hive_files
        self.output_dir = output_dir
        self.is_cancelled = False

    def cancel(self):
        self.is_cancelled = True
        
    def run(self):
        try:
            total_hives = len(self.hive_files)
            results = {}

            for i, hive_file in enumerate(self.hive_files):
                if self.is_cancelled:
                    break

                hive_name = os.path.basename(hive_file)
                self.status_updated.emit(f"Đang phân tích {hive_name}...")

                result = self.run_recmd_analysis(hive_file)
                if result:
                    results[hive_file] = result

                progress = int((i + 1) / total_hives * 100)
                self.progress_updated.emit(progress)

            if not self.is_cancelled:
                self.analysis_completed.emit(results)

        except Exception as e:
            self.error_occurred.emit(str(e))
            
    def run_recmd_analysis(self, hive_file):
        """Chạy RECmd trên một file hive duy nhất."""
        try:
            hive_name = os.path.splitext(os.path.basename(hive_file))[0]
            output_csv = os.path.join(self.output_dir, f"{hive_name}_analysis.csv")

            cmd = [
                self.recmd_path,
                "-f", hive_file,
                "--bn", self.batch_file,
                "--csv", self.output_dir,
                "--csvf", f"{hive_name}_analysis.csv",
                "--nl"
            ]

            process = subprocess.Popen(
                cmd,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                creationflags=subprocess.CREATE_NO_WINDOW if os.name == "nt" else 0
            )

            stdout, stderr = process.communicate()

            if process.returncode == 0 and os.path.exists(output_csv):
                return self.parse_csv_results(output_csv)
            else:
                raise Exception(f"RECmd thất bại: {stderr}")

        except Exception as e:
            raise Exception(f"Lỗi phân tích {os.path.basename(hive_file)}: {str(e)}")

    def parse_csv_results(self, csv_file):
        """Phân tích kết quả CSV từ RECmd."""
        results = []
        try:
            with open(csv_file, "r", encoding="utf-8", errors="ignore") as f:
                reader = csv.DictReader(f)
                for row in reader:
                    results.append(row)
            return results
        except Exception as e:
            raise Exception(f"Lỗi đọc CSV {csv_file}: {str(e)}")

# ============= Dialog Chọn Registry Hives =============

class HiveSelectionDialog(QDialog):
    """Dialog để chọn registry hives muốn phân tích."""
    
    def __init__(self, registry_files, parent=None):
        super().__init__(parent)
        self.registry_files = registry_files
        self.selected_files = []
        self.setup_ui()
        
    def setup_ui(self):
        """Thiết lập giao diện dialog."""
        self.setWindowTitle("Chọn Registry Hives để phân tích")
        self.setFixedSize(600, 400)
        self.setModal(True)
        
        layout = QVBoxLayout()
        
        # Header
        header_label = QLabel(f"🔍 Tìm thấy {len(self.registry_files)} Registry Hive files")
        header_label.setStyleSheet("font-size: 14px; font-weight: bold; margin: 10px;")
        layout.addWidget(header_label)
        
        # Instruction
        instruction = QLabel("Chọn những file hive mà bạn muốn phân tích:")
        instruction.setStyleSheet("margin: 5px; color: #666;")
        layout.addWidget(instruction)
        
        # Scroll area cho danh sách checkboxes
        scroll_area = QScrollArea()
        scroll_widget = QWidget()
        scroll_layout = QVBoxLayout(scroll_widget)
        
        self.checkboxes = []
        
        # Tạo checkbox cho từng hive file
        for file_path in self.registry_files:
            hive_type = self.detect_hive_type(str(file_path))
            file_name = file_path.name
            file_size = file_path.stat().st_size if file_path.exists() else 0
            
            # Tạo frame cho mỗi hive
            hive_frame = QFrame()
            hive_frame.setFrameStyle(QFrame.Box)
            hive_frame.setStyleSheet("margin: 2px; padding: 5px; background-color: #f9f9f9;")
            
            hive_layout = QVBoxLayout(hive_frame)
            
            # Checkbox chính
            checkbox = QCheckBox(f"📁 {file_name}")
            checkbox.setStyleSheet("font-weight: bold; font-size: 12px;")
            checkbox.file_path = file_path
            
            # Mặc định chọn các hive quan trọng
            if hive_type in ["SYSTEM", "SOFTWARE", "SAM", "SECURITY"]:
                checkbox.setChecked(True)
            
            # Thông tin chi tiết
            info_label = QLabel(f"   🏷️ Loại: {hive_type} | 📏 Kích thước: {file_size:,} bytes")
            info_label.setStyleSheet("color: #666; font-size: 10px; margin-left: 20px;")
            
            path_label = QLabel(f"   📂 Đường dẫn: {str(file_path)}")
            path_label.setStyleSheet("color: #888; font-size: 9px; margin-left: 20px;")
            path_label.setWordWrap(True)
            
            hive_layout.addWidget(checkbox)
            hive_layout.addWidget(info_label)
            hive_layout.addWidget(path_label)
            
            scroll_layout.addWidget(hive_frame)
            self.checkboxes.append(checkbox)
        
        scroll_area.setWidget(scroll_widget)
        scroll_area.setWidgetResizable(True)
        layout.addWidget(scroll_area)
        
        # Buttons
        button_layout = QHBoxLayout()
        
        select_all_btn = QPushButton("✅ Chọn tất cả")
        select_all_btn.clicked.connect(self.select_all)
        
        deselect_all_btn = QPushButton("❌ Bỏ chọn tất cả")
        deselect_all_btn.clicked.connect(self.deselect_all)
        
        select_important_btn = QPushButton("⭐ Chọn hive quan trọng")
        select_important_btn.clicked.connect(self.select_important)
        
        button_layout.addWidget(select_all_btn)
        button_layout.addWidget(deselect_all_btn)
        button_layout.addWidget(select_important_btn)
        button_layout.addStretch()
        
        layout.addLayout(button_layout)
        
        # OK/Cancel buttons
        ok_cancel_layout = QHBoxLayout()
        
        ok_btn = QPushButton("🚀 Bắt đầu phân tích")
        ok_btn.setStyleSheet("background-color: #4CAF50; color: white; font-weight: bold; padding: 8px;")
        ok_btn.clicked.connect(self.accept)
        
        cancel_btn = QPushButton("❌ Hủy")
        cancel_btn.clicked.connect(self.reject)
        
        ok_cancel_layout.addStretch()
        ok_cancel_layout.addWidget(cancel_btn)
        ok_cancel_layout.addWidget(ok_btn)
        
        layout.addLayout(ok_cancel_layout)
        
        self.setLayout(layout)
        
    def detect_hive_type(self, file_path):
        """Phát hiện loại hive từ tên file."""
        filename = os.path.basename(file_path).upper()
        
        if "SYSTEM" in filename:
            return "SYSTEM"
        elif "SOFTWARE" in filename:
            return "SOFTWARE"
        elif "SAM" in filename:
            return "SAM"
        elif "SECURITY" in filename:
            return "SECURITY"
        elif "NTUSER" in filename:
            return "NTUSER"
        elif "USRCLASS" in filename:
            return "USRCLASS"
        elif "DEFAULT" in filename:
            return "DEFAULT"
        else:
            return "UNKNOWN"
    
    def select_all(self):
        """Chọn tất cả checkboxes."""
        for checkbox in self.checkboxes:
            checkbox.setChecked(True)
    
    def deselect_all(self):
        """Bỏ chọn tất cả checkboxes."""
        for checkbox in self.checkboxes:
            checkbox.setChecked(False)
    
    def select_important(self):
        """Chọn chỉ các hive quan trọng."""
        important_types = ["SYSTEM", "SOFTWARE", "SAM", "SECURITY"]
        for checkbox in self.checkboxes:
            hive_type = self.detect_hive_type(str(checkbox.file_path))
            checkbox.setChecked(hive_type in important_types)
    
    def get_selected_files(self):
        """Lấy danh sách file được chọn."""
        selected = []
        for checkbox in self.checkboxes:
            if checkbox.isChecked():
                selected.append(checkbox.file_path)
        return selected

# ============= Widget Phân Tích Registry Chính =============

class RegistryAnalysis(QWidget):
    """Widget Phân Tích Registry - phù hợp với UI mới."""

    def __init__(self, main_window=None):
        super().__init__()
        self.main_window = main_window
        self.ui = Ui_RegistryAnalysisWidget()
        self.ui.setupUi(self)

        # Khởi tạo đường dẫn
        self._initialize_paths()

        # Quản lý trạng thái
        self.loaded_hives = {}
        self.analysis_results = {}
        self.bookmarks = []
        self.timeline_events = []
        self.current_case_id = None
        self.current_analysis_thread = None
        self.registry_objects = {}  # Cache các đối tượng registry
        self.reports_dir = None  # Thư mục lưu báo cáo

        # Models cho QTreeView và QTableView
        self.tree_model = QStandardItemModel()
        self.table_model = QStandardItemModel()

        # Thiết lập UI
        self.setup_ui()
        self.setup_connections()
        self.setup_quick_access()

        # Tải dữ liệu case nếu có sẵn (load ngầm)
        if main_window and hasattr(main_window, 'current_case_id'):
            # Sử dụng QTimer để load case sau khi UI đã hiển thị
            QTimer.singleShot(100, lambda: self.load_case_data(main_window.current_case_id))
            
    def _initialize_paths(self):
        """Khởi tạo đường dẫn công cụ."""
        try:
            from utils.path_utils import get_tools_dir
            self.tools_dir = get_tools_dir()
            self.recmd_path = os.path.join(self.tools_dir, "RECmd", "RECmd.exe")
            self.batch_dir = os.path.join(self.tools_dir, "RECmd", "BatchExamples")

            # Không khởi tạo output_dir ở đây - sẽ được thiết lập trong load_case_data
            self.output_dir = None
        except:
            # Đường dẫn dự phòng
            base_dir = os.path.dirname(os.path.dirname(os.path.dirname(__file__)))
            self.tools_dir = os.path.join(base_dir, "tools")
            self.recmd_path = os.path.join(self.tools_dir, "RECmd", "RECmd.exe")
            self.batch_dir = os.path.join(self.tools_dir, "RECmd", "BatchExamples")

            # Không khởi tạo output_dir ở đây - sẽ được thiết lập trong load_case_data
            self.output_dir = None

    def setup_ui(self):
        """Thiết lập các thành phần UI."""
        # Thiết lập thuộc tính cửa sổ
        self.setWindowTitle("Phân Tích Registry - Công Cụ Điều Tra Pháp Y Số")

        # Cập nhật header
        self.update_case_info()
        self.update_status("Sẵn sàng")
        

        # Thiết lập models cho QTreeView
        self.tree_model.setHorizontalHeaderLabels(["Các Key Registry"])
        self.ui.registryTree.setModel(self.tree_model)
        self.ui.registryTree.setContextMenuPolicy(Qt.CustomContextMenu)

        # Thiết lập model cho QTableView
        self.ui.valuesTable.setModel(self.table_model)
        self.ui.valuesTable.setContextMenuPolicy(Qt.CustomContextMenu)

        # Cấu hình hex view
        self.ui.hexView.setReadOnly(True)
        font = QFont("Consolas", 9)
        self.ui.hexView.setFont(font)

        # Cấu hình decoded view
        self.ui.decodedView.setReadOnly(True)

        # Thiết lập bảng timeline
        self.setup_timeline_table()

        # Thiết lập kích thước splitter ban đầu
        self.ui.mainSplitter.setSizes([350, 850])
        self.ui.verticalSplitter.setSizes([400, 200])
        
    def setup_connections(self):
        """Kết nối signals và slots."""
        # Các hành động toolbar
        self.ui.btnLoadHive.clicked.connect(self.load_registry_hives)
        self.ui.txtSearch.textChanged.connect(self.on_search_text_changed)
        self.ui.txtSearch.returnPressed.connect(self.perform_search)
        self.ui.btnSearchOptions.clicked.connect(self.show_search_menu)
        self.ui.btnExport.clicked.connect(self.show_export_menu)
        

        # Các hành động tree
        self.ui.registryTree.clicked.connect(self.on_tree_item_clicked)
        self.ui.registryTree.customContextMenuRequested.connect(self.show_tree_context_menu)
        self.ui.btnExpandAll.clicked.connect(lambda: self.ui.registryTree.expandAll())
        self.ui.btnCollapseAll.clicked.connect(lambda: self.ui.registryTree.collapseAll())
        self.ui.txtTreeFilter.textChanged.connect(self.filter_tree)

        # Lựa chọn bảng
        self.ui.valuesTable.selectionModel().selectionChanged.connect(self.on_value_selected)

        # Bookmarks
        self.ui.btnAddBookmark.clicked.connect(self.add_bookmark)
        self.ui.btnRemoveBookmark.clicked.connect(self.remove_bookmark)
        self.ui.btnGoToBookmark.clicked.connect(self.go_to_bookmark)
        self.ui.bookmarksList.itemDoubleClicked.connect(self.bookmark_double_clicked)

        # Truy cập nhanh (nếu có trong UI)
        # self.ui.quickAccessList.itemDoubleClicked.connect(self.quick_access_double_clicked)

        # Combo định dạng
        self.ui.cmbFormat.currentTextChanged.connect(self.update_decoded_view)

        # Ghi chú
        self.ui.btnSaveNotes.clicked.connect(self.save_notes)

        # Thanh đường dẫn
        self.ui.btnCopyPath.clicked.connect(self.copy_current_path)
        
        
    def setup_quick_access(self):
        """Thiết lập các vị trí truy cập nhanh - đã được điền sẵn trong file UI."""
        pass  # Các mục truy cập nhanh đã được thiết lập trong file UI

    def setup_timeline_table(self):
        """Thiết lập các cột của bảng timeline."""
        self.ui.timelineTable.setColumnCount(4)
        self.ui.timelineTable.setHorizontalHeaderLabels(
            ["Thời gian", "Key", "Hành động", "Chi tiết"]
        )
        self.ui.timelineTable.horizontalHeader().setStretchLastSection(True)
        self.ui.timelineTable.setAlternatingRowColors(True)
        self.ui.timelineTable.setSortingEnabled(True)
        
    def load_case_data(self, case_id):
        """Tải dữ liệu cụ thể của case."""
            
        self.current_case_id = case_id
        self.update_case_info()

        # Thử tự động tải file registry từ case và thiết lập thư mục output
        try:
            from models.db_manager import DatabaseManager
            db = DatabaseManager()
            db.connect()

            case_info = db.get_case_with_investigator(case_id)
            if case_info and case_info.get('archive_path'):
                # Thiết lập thư mục output và reports trong thư mục case
                case_path = case_info['archive_path']
                self.output_dir = os.path.join(case_path, "analysis_results", "registry")
                self.reports_dir = os.path.join(case_path, "reports", "registry")
                
                # Tạo các thư mục cần thiết
                os.makedirs(self.output_dir, exist_ok=True)
                os.makedirs(self.reports_dir, exist_ok=True)

                # Tải kết quả phân tích đã có (nếu có) - PHẢI TRƯỚC auto_load_case_registry
                print(f"🔄 Đang tải kết quả cũ cho case {self.current_case_id}")
                self._load_existing_analysis_results(db)

                # Delay một chút để đảm bảo restore hoàn tất trước khi load registry files
                QTimer.singleShot(500, lambda: self.auto_load_case_registry(case_path))
                self.update_status(f"Case đã tải - Output: {self.output_dir}", "green")
            else:
                # Dự phòng về temp nếu không có đường dẫn case
                self._set_fallback_output_dir()
            
            db.disconnect()
        except Exception as e:
            print(f"Lỗi tải dữ liệu case: {e}")
            # Dự phòng về temp nếu có lỗi
            self._set_fallback_output_dir()

    def _set_fallback_output_dir(self):
        """Thiết lập thư mục output dự phòng khi không có đường dẫn case."""
        try:
            from utils.path_utils import get_temp_dir
            temp_root = get_temp_dir() if callable(get_temp_dir) else "temp"
            self.output_dir = os.path.join(temp_root, "registry_analysis")
            self.reports_dir = os.path.join(temp_root, "reports", "registry")
        except:
            base_dir = os.path.dirname(os.path.dirname(os.path.dirname(__file__)))
            self.output_dir = os.path.join(base_dir, "temp", "registry_analysis")
            self.reports_dir = os.path.join(base_dir, "temp", "reports", "registry")

        os.makedirs(self.output_dir, exist_ok=True)
        os.makedirs(self.reports_dir, exist_ok=True)
        self.update_status("Sử dụng thư mục output tạm thời", "yellow")

    def _load_existing_analysis_results(self, db):
        """Tải kết quả phân tích đã có từ database và file CSV."""
        try:
            print(f"🔍 Đang kiểm tra kết quả cũ cho case: {self.current_case_id}")
            
            # Lấy kết quả phân tích registry cho case hiện tại
            results = db.get_results_by_case(self.current_case_id)
            registry_results = [r for r in results if 'registry' in r.get('tool_used', '').lower()]
            
            print(f"📊 Tìm thấy {len(results)} kết quả tổng cộng")
            print(f"🔧 Tìm thấy {len(registry_results)} kết quả registry")
            
            if registry_results:
                print("✅ Có kết quả cũ, đang restore từ CSV...")
                self.update_status(f"Đã tìm thấy {len(registry_results)} kết quả cũ", "green")
                
                # Thử load lại kết quả từ file CSV
                self._restore_analysis_results_from_files()
                
                # Tải timeline từ database nếu có
                self._load_timeline_from_database(db)
                
            else:
                print("❌ Không có kết quả cũ")
                self.update_status("Chưa có kết quả phân tích nào được lưu", "yellow")
                
        except Exception as e:
            print(f"❌ Lỗi tải kết quả đã có: {e}")
            import traceback
            traceback.print_exc()
            
    def _restore_analysis_results_from_files(self):
        """Khôi phục kết quả phân tích từ file CSV đã lưu."""
        print(f"🔍 Kiểm tra output_dir: {self.output_dir}")
        
        if not self.output_dir or not os.path.exists(self.output_dir):
            print(f"❌ Output dir không tồn tại: {self.output_dir}")
            return
            
        # Liệt kê tất cả file trong output_dir
        try:
            all_files = os.listdir(self.output_dir)
            print(f"📁 Tất cả files trong output_dir: {all_files}")
        except Exception as e:
            print(f"❌ Lỗi list files: {e}")
            
        try:
            import glob
            csv_pattern = os.path.join(self.output_dir, "*_analysis.csv")
            csv_files = glob.glob(csv_pattern)
            print(f"🔍 Tìm CSV với pattern: {csv_pattern}")
            print(f"📁 Tìm thấy {len(csv_files)} CSV files: {csv_files}")
            
            if csv_files:
                restored_results = {}
                total_records = 0
                
                for csv_file in csv_files:
                    try:
                        print(f"📄 Đang xử lý CSV: {csv_file}")
                        
                        # Xác định hive file từ tên CSV
                        csv_name = os.path.basename(csv_file)
                        hive_name = csv_name.replace('_analysis.csv', '')
                        print(f"🏷️ Hive name từ CSV: {hive_name}")
                        
                        # Tìm hive file tương ứng trong case
                        possible_hive_paths = self._find_hive_file_by_name(hive_name)
                        print(f"📂 Tìm thấy {len(possible_hive_paths)} hive paths: {possible_hive_paths}")
                        
                        if possible_hive_paths:
                            hive_file = possible_hive_paths[0]  # Lấy file đầu tiên
                            
                            # Parse CSV results
                            csv_results = self._parse_csv_file(csv_file)
                            if csv_results:
                                restored_results[hive_file] = csv_results
                                total_records += len(csv_results)
                                print(f"✅ Restored {len(csv_results)} records từ {csv_file}")
                        else:
                            print(f"❌ Không tìm thấy hive file cho {hive_name}")
                                
                    except Exception as e:
                        print(f"❌ Lỗi load CSV {csv_file}: {e}")
                
                if restored_results:
                    self.analysis_results = restored_results
                    print(f"🎉 Tổng cộng restored {total_records} artifacts từ {len(restored_results)} hives")
                    self.update_status(f"Đã khôi phục {total_records} artifacts từ {len(csv_files)} file", "green")
                else:
                    print("❌ Không restore được kết quả nào")
            else:
                print("❌ Không tìm thấy CSV file nào")
                    
        except Exception as e:
            print(f"❌ Lỗi khôi phục kết quả: {e}")
            import traceback
            traceback.print_exc()
            
    def _find_hive_file_by_name(self, hive_name):
        """Tìm file hive theo tên trong case hiện tại."""
        try:
            from models.db_manager import DatabaseManager
            db = DatabaseManager()
            db.connect()
            
            case_info = db.get_case_with_investigator(self.current_case_id)
            db.disconnect()
            
            if case_info and case_info.get('archive_path'):
                archive_path = Path(case_info['archive_path'])
                
                # Tìm file có tên chứa hive_name
                patterns = [
                    f"**/*{hive_name}*",
                    f"**/{hive_name}",
                    f"**/{hive_name.upper()}*",
                    f"**/{hive_name.lower()}*"
                ]
                
                found_files = []
                for pattern in patterns:
                    found_files.extend(archive_path.glob(pattern))
                
                return [str(f) for f in found_files if f.is_file()]
                
        except Exception as e:
            print(f"Lỗi tìm hive file: {e}")
            
        return []
        
    def _parse_csv_file(self, csv_file):
        """Parse file CSV và trả về kết quả."""
        try:
            results = []
            with open(csv_file, "r", encoding="utf-8", errors="ignore") as f:
                reader = csv.DictReader(f)
                for row in reader:
                    results.append(row)
            return results
        except Exception as e:
            print(f"Lỗi parse CSV {csv_file}: {e}")
            return []
            
    def _load_timeline_from_database(self, db):
        """Tải timeline từ activity logs trong database."""
        try:
            activity_logs = db.get_activity_logs(case_id=self.current_case_id)
            registry_logs = [log for log in activity_logs if 'registry' in log.get('action', '').lower()]
            
            if registry_logs:
                # Cập nhật timeline table
                self.ui.timelineTable.setRowCount(len(registry_logs))
                
                for i, log in enumerate(registry_logs):
                    self.ui.timelineTable.setItem(i, 0, QTableWidgetItem(str(log.get('timestamp', ''))))
                    self.ui.timelineTable.setItem(i, 1, QTableWidgetItem('Database Activity'))
                    self.ui.timelineTable.setItem(i, 2, QTableWidgetItem(log.get('action', '')))
                    self.ui.timelineTable.setItem(i, 3, QTableWidgetItem(log.get('details', '')))
                
                self.ui.timelineTable.resizeColumnsToContents()
                
        except Exception as e:
            print(f"Lỗi tải timeline từ database: {e}")

    def _check_and_run_analysis_if_needed(self):
        """Kiểm tra và chỉ chạy phân tích nếu cần thiết."""
        if not self.current_case_id:
            # Không có case, chạy phân tích luôn
            self.start_comprehensive_analysis()
            return

        try:
            from models.db_manager import DatabaseManager
            db = DatabaseManager()
            db.connect()

            # Kiểm tra xem đã có kết quả phân tích registry chưa
            results = db.get_results_by_case(self.current_case_id)
            registry_results = [r for r in results if 'registry' in r.get('tool_used', '').lower()]

            db.disconnect()

            if registry_results:
                # Đã có kết quả cũ, hỏi user có muốn chạy lại không
                reply = QMessageBox.question(
                    self,
                    "Đã có kết quả phân tích",
                    f"Case này đã có {len(registry_results)} kết quả phân tích Registry.\n\n"
                    f"Bạn có muốn chạy phân tích mới để cập nhật kết quả?",
                    QMessageBox.Yes | QMessageBox.No,
                    QMessageBox.No  # Default là No
                )
                
                if reply == QMessageBox.Yes:
                    self.start_comprehensive_analysis()
                else:
                    self.update_status("Sử dụng kết quả phân tích đã có", "green")
            else:
                # Chưa có kết quả, chạy phân tích tự động
                self.start_comprehensive_analysis()

        except Exception as e:
            print(f"Lỗi kiểm tra kết quả cũ: {e}")
            # Nếu có lỗi, chạy phân tích luôn
            self.start_comprehensive_analysis()

    def _auto_load_registry_files_if_needed(self, registry_files):
        """Tự động load registry files nếu cần thiết."""
        if not self.current_case_id:
            # Không có case, load file đầu tiên
            if registry_files:
                self.process_hive_files([str(registry_files[0])])
            return

        try:
            from models.db_manager import DatabaseManager
            db = DatabaseManager()
            db.connect()

            # Kiểm tra xem đã có kết quả phân tích registry chưa
            results = db.get_results_by_case(self.current_case_id)
            registry_results = [r for r in results if 'registry' in r.get('tool_used', '').lower()]

            db.disconnect()

            if not registry_results:
                # Chưa có kết quả, tự động load file đầu tiên
                if registry_files:
                    self.process_hive_files([str(registry_files[0])])
            else:
                # Đã có kết quả, chỉ load registry files mà KHÔNG chạy phân tích
                self.update_status("Đã có kết quả phân tích registry từ trước", "green")
                if registry_files:
                    # Load hive files để hiển thị tree nhưng không chạy analysis
                    self._load_hive_files_without_analysis([str(registry_files[0])])

        except Exception as e:
            print(f"Lỗi kiểm tra auto load: {e}")
            # Nếu có lỗi, load file đầu tiên
            if registry_files:
                self.process_hive_files([str(registry_files[0])])

    def _load_hive_files_without_analysis(self, file_paths):
        """Load registry files chỉ để hiển thị tree, không chạy phân tích."""
        if not REGISTRY_AVAILABLE:
            return
            
        valid_hives = []
        
        for file_path in file_paths:
            try:
                # Parse registry file
                registry = Registry.Registry(file_path)
                self.registry_objects[file_path] = registry
                
                hive_info = {
                    "path": file_path,
                    "name": os.path.basename(file_path),
                    "type": self.detect_hive_type(file_path),
                    "size": os.path.getsize(file_path),
                    "modified": datetime.fromtimestamp(os.path.getmtime(file_path)),
                    "registry": registry
                }
                
                valid_hives.append(hive_info)
                self.loaded_hives[file_path] = hive_info
                
            except Exception as e:
                print(f"Error loading {file_path}: {e}")
                
        if valid_hives:
            self.build_registry_tree()
            self.update_timeline()
            self.update_status(f"Đã tải {len(valid_hives)} hive(s) (sử dụng kết quả cũ)", "green")
            # KHÔNG gọi _check_and_run_analysis_if_needed() ở đây
            
    def update_case_info(self):
        """Cập nhật thông tin case trong header."""
        if self.current_case_id:
            self.ui.caseInfoLabel.setText(f"Case ID: {self.current_case_id}")
        else:
            self.ui.caseInfoLabel.setText("Case: Chưa chọn")

    def update_status(self, status, color="green"):
        """Cập nhật chỉ báo trạng thái."""
        color_map = {
            "green": "#90EE90",
            "yellow": "#FFD700",
            "red": "#FF6B6B"
        }
        self.ui.statusIndicator.setText(f"● {status}")
        self.ui.statusIndicator.setStyleSheet(f"color: {color_map.get(color, '#90EE90')}; font-size: 12px;")

        # Cũng cập nhật status bar
        self.ui.statusBar.showMessage(status, 5000)
            
    def auto_load_case_registry(self, archive_path):
        """Tự động tải file registry từ kho lưu trữ case."""
        archive_path = Path(archive_path)
        if not archive_path.exists():
            return

        self.update_status("Đang tìm kiếm registry hives...", "yellow")
        QTimer.singleShot(200, lambda: self._search_registry_files(archive_path))
        
    def _search_registry_files(self, archive_path):
        """Tìm kiếm registry files trong background."""
        # Các pattern file registry thường gặp
        patterns = [
            "**/*SYSTEM", "**/*SOFTWARE", "**/*SAM",
            "**/*SECURITY", "**/*NTUSER.DAT", "**/*UsrClass.dat"
        ]

        registry_files = []
        for pattern in patterns:
            found_files = list(archive_path.glob(pattern))
            registry_files.extend(found_files)

        if registry_files:
            # Kiểm tra xem đã có kết quả phân tích chưa
            print(f"🔍 Kiểm tra analysis_results: {len(self.analysis_results) if self.analysis_results else 0} items")
            print(f"📊 Analysis results keys: {list(self.analysis_results.keys()) if self.analysis_results else 'None'}")
            
            # Sử dụng hàm kiểm tra mạnh mẽ hơn
            has_results = self.has_existing_results()
            print(f"🔍 has_existing_results() = {has_results}")
            
            if has_results and self.analysis_results:
                print("✅ Đã có kết quả cũ, load registry files mà không hiển thị dialog")
                # Đã có kết quả cũ, chỉ load registry files mà không hiển thị dialog
                self._load_registry_files_without_dialog(registry_files)
            else:
                print("❌ Chưa có kết quả hoặc chưa restore được, hiển thị dialog để user chọn hives")
                # Chưa có kết quả, hiển thị dialog để user chọn hives muốn phân tích
                self.show_hive_selection_dialog(registry_files)
        else:
            self.update_status("Không tìm thấy registry file nào trong case", "yellow")
    
    def show_hive_selection_dialog(self, registry_files):
        """Hiển thị dialog để chọn hives muốn phân tích."""
        dialog = HiveSelectionDialog(registry_files, self)
        
        if dialog.exec_() == QDialog.Accepted:
            selected_files = dialog.get_selected_files()
            
            if selected_files:
                self.update_status(f"Đã chọn {len(selected_files)} hives", "green")
                self.process_hive_files([str(f) for f in selected_files])
                QTimer.singleShot(1000, self.start_silent_analysis)
            else:
                self.update_status("Không có hive nào được chọn", "yellow")
        else:
            self.update_status("Đã hủy chọn hives", "yellow")
    
    def has_existing_results(self):
        """Kiểm tra xem có kết quả cũ không (cho debug)."""
        if not self.current_case_id:
            return False
            
        try:
            from models.db_manager import DatabaseManager
            db = DatabaseManager()
            db.connect()
            
            results = db.get_results_by_case(self.current_case_id)
            registry_results = [r for r in results if 'registry' in r.get('tool_used', '').lower()]
            
            db.disconnect()
            
            # Kiểm tra cả database và CSV files
            has_db_results = len(registry_results) > 0
            has_csv_files = False
            
            if self.output_dir and os.path.exists(self.output_dir):
                import glob
                csv_files = glob.glob(os.path.join(self.output_dir, "*_analysis.csv"))
                has_csv_files = len(csv_files) > 0
            
            print(f"🔍 has_existing_results: DB={has_db_results}, CSV={has_csv_files}")
            return has_db_results and has_csv_files
            
        except Exception as e:
            print(f"❌ Lỗi kiểm tra kết quả cũ: {e}")
            return False
    
    def _load_registry_files_without_dialog(self, registry_files):
        """Load registry files mà không hiển thị dialog (dành cho kết quả cũ)."""
        if not REGISTRY_AVAILABLE:
            return
            
        # Load tất cả hive files tìm thấy
        file_paths = [str(f) for f in registry_files]
        self.process_hive_files(file_paths)
        
        self.update_status(f"Đã load {len(registry_files)} hives (sử dụng kết quả cũ)", "green")
            
    def load_registry_hives(self):
        """Tải file registry hive."""
        file_dialog = QFileDialog()
        file_dialog.setFileMode(QFileDialog.ExistingFiles)
        file_dialog.setNameFilter("Registry Hives (*);;Tất cả Files (*)")
        file_dialog.setWindowTitle("Chọn File Registry Hive")

        if file_dialog.exec_():
            files = file_dialog.selectedFiles()
            if files:
                # Convert string paths to Path objects
                path_objects = [Path(f) for f in files]
                # Hiển thị dialog chọn hive
                self.show_hive_selection_dialog(path_objects)
                
    def process_hive_files(self, file_paths):
        """Xử lý và tải file hive."""
        if not REGISTRY_AVAILABLE:
            QMessageBox.warning(
                self,
                "Thiếu Thư Viện",
                "Thư viện python-registry chưa được cài đặt.\n"
                "Vui lòng cài đặt bằng: pip install python-registry"
            )
            return
            
        valid_hives = []
        
        for file_path in file_paths:
            try:
                # Parse registry file
                registry = Registry.Registry(file_path)
                self.registry_objects[file_path] = registry
                
                hive_info = {
                    "path": file_path,
                    "name": os.path.basename(file_path),
                    "type": self.detect_hive_type(file_path),
                    "size": os.path.getsize(file_path),
                    "modified": datetime.fromtimestamp(os.path.getmtime(file_path)),
                    "registry": registry
                }
                
                valid_hives.append(hive_info)
                self.loaded_hives[file_path] = hive_info
                
            except Exception as e:
                print(f"Error loading {file_path}: {e}")
                
        if valid_hives:
            self.build_registry_tree()
            self.update_timeline()
            self.update_status(f"Đã tải {len(valid_hives)} hive(s)", "green")
            
    def detect_hive_type(self, file_path):
        """Phát hiện loại hive từ tên file."""
        filename = os.path.basename(file_path).upper()

        if "SYSTEM" in filename:
            return "SYSTEM"
        elif "SOFTWARE" in filename:
            return "SOFTWARE"
        elif "SAM" in filename:
            return "SAM"
        elif "SECURITY" in filename:
            return "SECURITY"
        elif "NTUSER" in filename:
            return "NTUSER"
        elif "USRCLASS" in filename:
            return "USRCLASS"
        elif "DEFAULT" in filename:
            return "DEFAULT"
        else:
            return "UNKNOWN"
            
    def build_registry_tree(self):
        """Xây dựng tree view registry với QStandardItemModel."""
        self.tree_model.clear()
        self.tree_model.setHorizontalHeaderLabels(["Các Key Registry"])

        for hive_path, hive_info in self.loaded_hives.items():
            # Tạo hive root item
            hive_item = QStandardItem(f"{hive_info['name']} ({hive_info['type']})")
            hive_item.setData({
                "type": "hive",
                "path": hive_path,
                "info": hive_info
            }, Qt.UserRole)

            self.tree_model.appendRow(hive_item)

            # Thêm registry keys
            try:
                registry = hive_info['registry']
                root = registry.root()
                self.add_key_to_tree(hive_item, root, registry)
            except Exception as e:
                error_item = QStandardItem(f"Lỗi: {str(e)}")
                hive_item.appendRow(error_item)

        # Mở rộng level đầu tiên
        self.ui.registryTree.expandToDepth(0)
        
    def add_key_to_tree(self, parent_item, key, registry, depth=0, max_depth=50):
        """Đệ quy thêm registry keys vào tree."""
        if depth >= max_depth:
            return

        try:
            # Tạo item cho key này
            key_name = key.name() if key.name() else "Root"
            key_item = QStandardItem(key_name)
            key_item.setData({
                "type": "key",
                "key": key,
                "path": key.path(),
                "registry": registry
            }, Qt.UserRole)

            parent_item.appendRow(key_item)

            # Thêm subkeys
            for subkey in key.subkeys():
                self.add_key_to_tree(key_item, subkey, registry, depth + 1, max_depth)

        except Exception as e:
            pass  # Im lặng bỏ qua các key có vấn đề
            
    def on_tree_item_clicked(self, index):
        """Xử lý sự kiện click item trong tree."""
        item = self.tree_model.itemFromIndex(index)
        if not item:
            return

        data = item.data(Qt.UserRole)
        if not data:
            return

        if data["type"] == "key":
            # Hiển thị giá trị key
            self.show_key_values(data["key"])

            # Cập nhật đường dẫn
            self.ui.txtCurrentPath.setText(data["path"])

            # Cập nhật phân tích
            self.analyze_key(data["key"], data["path"])

        elif data["type"] == "hive":
            # Hiển thị thông tin hive
            self.show_hive_info(data["info"])
            
    def show_key_values(self, key):
        """Hiển thị giá trị của key được chọn trong QTableView."""
        # Xóa table model
        self.table_model.clear()
        self.table_model.setHorizontalHeaderLabels(["Tên", "Kiểu", "Dữ liệu"])

        try:
            values = []
            for value in key.values():
                values.append({
                    "name": value.name() if value.name() else "(Mặc định)",
                    "type": value.value_type_str(),
                    "data": self.format_value_data(value),
                    "raw_data": value.raw_data()
                })

            # Điền dữ liệu vào model
            for val in values:
                row = []
                name_item = QStandardItem(val["name"])
                name_item.setData(val["raw_data"], Qt.UserRole)  # Lưu raw data
                row.append(name_item)
                row.append(QStandardItem(val["type"]))
                row.append(QStandardItem(str(val["data"])))
                self.table_model.appendRow(row)

            # Tự động điều chỉnh kích thước cột
            self.ui.valuesTable.resizeColumnsToContents()

        except Exception as e:
            print(f"Lỗi hiển thị giá trị: {e}")
            
    def format_value_data(self, value):
        """Định dạng dữ liệu value để hiển thị."""
        try:
            data = value.value()
            if isinstance(data, str):
                return data.replace('\x00', '').strip()
            elif isinstance(data, (int, float)):
                return str(data)
            elif isinstance(data, bytes):
                try:
                    return data.decode('utf-16le', errors='ignore').rstrip('\x00')
                except:
                    return f"<Nhị phân: {len(data)} bytes>"
            elif isinstance(data, list):
                return "; ".join(str(item) for item in data[:5])
            else:
                return str(data)
        except:
            return "<Lỗi đọc giá trị>"
            
    def on_value_selected(self, selected, deselected):
        """Xử lý sự kiện lựa chọn value."""
        indexes = selected.indexes()
        if not indexes:
            return

        # Lấy cột đầu tiên (tên) của hàng được chọn
        row = indexes[0].row()
        name_item = self.table_model.item(row, 0)

        if name_item:
            # Lấy raw data được lưu trong item
            raw_data = name_item.data(Qt.UserRole)

            if raw_data:
                # Cập nhật hex view
                self.ui.hexView.setPlainText(format_as_hex(raw_data))

                # Cập nhật decoded view
                self.update_decoded_view()
            
    def update_decoded_view(self):
        """Cập nhật view dữ liệu đã giải mã."""
        # Lấy hàng được chọn
        indexes = self.ui.valuesTable.selectionModel().selectedIndexes()
        if not indexes:
            return

        row = indexes[0].row()
        name_item = self.table_model.item(row, 0)

        if name_item:
            raw_data = name_item.data(Qt.UserRole)

            if raw_data:
                format_type = self.ui.cmbFormat.currentText()
                decoded = decode_registry_data(raw_data, format_type)
                self.ui.decodedView.setPlainText(decoded)
            
    def analyze_key(self, key, path):
        """Phân tích registry key để tìm forensic artifacts."""
        analysis_text = f"Phân Tích Registry Key\n"
        analysis_text += f"{'='*50}\n"
        analysis_text += f"Đường dẫn: {path}\n"
        analysis_text += f"Sửa đổi lần cuối: {key.timestamp()}\n\n"

        # Kiểm tra forensic artifacts đã biết
        path_upper = path.upper()

        if "USERASSIST" in path_upper:
            analysis_text += "📌 Phát hiện Key UserAssist\n"
            analysis_text += "Key này chứa lịch sử thực thi chương trình.\n"
            analysis_text += "Giá trị được mã hóa ROT13.\n"

        elif "RUN" in path_upper and "RUNONCE" not in path_upper:
            analysis_text += "🚀 Phát hiện Entry Khởi động tự động\n"
            analysis_text += "Các chương trình ở đây chạy khi user đăng nhập.\n"

        elif "SHELLBAGS" in path_upper:
            analysis_text += "📁 Phát hiện Shellbags\n"
            analysis_text += "Chứa lịch sử truy cập thư mục và tùy chỉnh.\n"

        elif "TYPEDURLS" in path_upper:
            analysis_text += "🌐 Phát hiện Typed URLs\n"
            analysis_text += "Chứa URLs được nhập trong Internet Explorer/Edge.\n"

        elif "MOUNTEDDEVICES" in path_upper:
            analysis_text += "💾 Phát hiện Mounted Devices\n"
            analysis_text += "Hiển thị lịch sử thiết bị lưu trữ đã kết nối.\n"

        elif "USBSTOR" in path_upper:
            analysis_text += "🔌 Lịch sử thiết bị USB\n"
            analysis_text += "Chứa thông tin về thiết bị USB đã kết nối.\n"

        self.ui.analysisView.setHtml(f"<pre>{analysis_text}</pre>")
        
    def update_timeline(self):
        """Cập nhật timeline với các thay đổi registry."""
        self.ui.timelineTable.setRowCount(0)
        timeline_events = []

        for hive_path, hive_info in self.loaded_hives.items():
            try:
                registry = hive_info['registry']
                self.collect_timeline_events(registry.root(), timeline_events, hive_info['name'])
            except:
                pass

        # Sắp xếp theo timestamp
        timeline_events.sort(key=lambda x: x['timestamp'], reverse=True)

        # Thêm vào bảng (giới hạn 100 sự kiện gần nhất)
        self.ui.timelineTable.setRowCount(min(len(timeline_events), 100))

        for i, event in enumerate(timeline_events[:100]):
            self.ui.timelineTable.setItem(i, 0, QTableWidgetItem(event['timestamp'].strftime('%Y-%m-%d %H:%M:%S')))
            self.ui.timelineTable.setItem(i, 1, QTableWidgetItem(event['key']))
            self.ui.timelineTable.setItem(i, 2, QTableWidgetItem(event['action']))
            self.ui.timelineTable.setItem(i, 3, QTableWidgetItem(event['details']))

        self.ui.timelineTable.resizeColumnsToContents()
        
    def collect_timeline_events(self, key, events, hive_name, depth=0, max_depth=3):
        """Thu thập timeline events từ registry keys."""
        if depth >= max_depth:
            return

        try:
            if key.timestamp():
                events.append({
                    'timestamp': key.timestamp(),
                    'key': key.path(),
                    'action': 'Đã sửa đổi',
                    'details': f'Hive: {hive_name}'
                })

            for subkey in key.subkeys():
                self.collect_timeline_events(subkey, events, hive_name, depth + 1, max_depth)
        except:
            pass
            
    def add_bookmark(self):
        """Thêm vị trí hiện tại vào bookmarks."""
        current_index = self.ui.registryTree.currentIndex()
        if not current_index.isValid():
            return

        item = self.tree_model.itemFromIndex(current_index)
        if not item:
            return

        data = item.data(Qt.UserRole)
        if data and data["type"] == "key":
            bookmark_text = data["path"]

            # Kiểm tra xem đã bookmark chưa
            for i in range(self.ui.bookmarksList.count()):
                if self.ui.bookmarksList.item(i).text() == bookmark_text:
                    return

            # Thêm bookmark
            list_item = QListWidgetItem(bookmark_text)
            list_item.setData(Qt.UserRole, data)
            self.ui.bookmarksList.addItem(list_item)
            self.bookmarks.append(data)

            self.update_status("Đã thêm bookmark", "green")
            
    def remove_bookmark(self):
        """Xóa bookmark đã chọn."""
        current = self.ui.bookmarksList.currentItem()
        if current:
            row = self.ui.bookmarksList.row(current)
            self.ui.bookmarksList.takeItem(row)
            if row < len(self.bookmarks):
                del self.bookmarks[row]
                
    def go_to_bookmark(self):
        """Điều hướng đến bookmark đã chọn."""
        current = self.ui.bookmarksList.currentItem()
        if current:
            data = current.data(Qt.UserRole)
            if data:
                # Tìm và chọn item trong tree
                self.find_and_select_tree_item(data["path"])
                
    def bookmark_double_clicked(self, item):
        """Xử lý sự kiện double-click bookmark."""
        self.go_to_bookmark()

                
    def find_and_select_tree_item(self, target_path):
        """Tìm và chọn item trong tree theo đường dẫn."""
        # Implementation sẽ tìm trong tree và chọn item phù hợp
        self.update_status(f"Điều hướng đến: {target_path}", "yellow")

    # def on_quick_load_selected(self, index):
    #     """Xử lý sự kiện lựa chọn quick load."""
    #     if index <= 0:
    #         return

    #     file_path = self.ui.cmbQuickLoad.itemData(index)
    #     if file_path and os.path.exists(file_path):
    #         self.process_hive_files([file_path])

    def on_search_text_changed(self, text):
        """Xử lý sự kiện thay đổi text tìm kiếm."""
        if len(text) >= 3:
            # Có thể implement live search
            pass
            
    def perform_search(self):
        """Thực hiện tìm kiếm registry."""
        search_text = self.ui.txtSearch.text()
        if not search_text:
            return

        self.update_status(f"Đang tìm kiếm: {search_text}", "yellow")

        # Implementation tìm kiếm sẽ được thêm vào đây
        QMessageBox.information(self, "Tìm kiếm", f"Chức năng tìm kiếm cho '{search_text}' đang được triển khai.")
        
    def filter_tree(self, text):
        """Lọc tree view."""
        # Implementation lọc tree
        pass

    def copy_current_path(self):
        """Sao chép đường dẫn hiện tại vào clipboard."""
        path = self.ui.txtCurrentPath.text()
        if path:
            QApplication.clipboard().setText(path)
            self.update_status("Đã sao chép đường dẫn vào clipboard", "green")

    def save_notes(self):
        """Lưu ghi chú điều tra."""
        notes = self.ui.notesEdit.toPlainText()
        if notes:
            # Lưu vào database hoặc file
            self.update_status("Đã lưu ghi chú", "green")
            
    def show_hive_info(self, hive_info):
        """Hiển thị thông tin hive."""
        info_text = f"Thông Tin Hive\n"
        info_text += f"{'='*50}\n"
        info_text += f"Tên: {hive_info['name']}\n"
        info_text += f"Loại: {hive_info['type']}\n"
        info_text += f"Đường dẫn: {hive_info['path']}\n"
        info_text += f"Kích thước: {hive_info['size']:,} bytes\n"
        info_text += f"Sửa đổi: {hive_info['modified']}\n"

        self.ui.analysisView.setHtml(f"<pre>{info_text}</pre>")
        
    def show_tree_context_menu(self, position):
        """Hiển thị context menu cho tree."""
        menu = QMenu()

        copy_path = menu.addAction("📋 Sao Chép Đường Dẫn")
        add_bookmark = menu.addAction("⭐ Thêm Bookmark")
        menu.addSeparator()
        export_key = menu.addAction("💾 Xuất Key")

        action = menu.exec_(self.ui.registryTree.mapToGlobal(position))

        if action == copy_path:
            self.copy_current_path()
        elif action == add_bookmark:
            self.add_bookmark()
            
    def show_search_menu(self):
        """Hiển thị menu tùy chọn tìm kiếm."""
        menu = QMenu()

        menu.addAction("🔍 Phân biệt hoa thường").setCheckable(True)
        menu.addAction("📝 Biểu thức chính quy").setCheckable(True)
        menu.addAction("🗑️ Bao gồm đã xóa").setCheckable(True)
        menu.addSeparator()
        menu.addAction("🐛 Debug: Kiểm tra kết quả cũ", self.debug_check_existing_results)

        menu.exec_(self.ui.btnSearchOptions.mapToGlobal(self.ui.btnSearchOptions.rect().bottomLeft()))
        
    def debug_check_existing_results(self):
        """Debug: Kiểm tra kết quả cũ."""
        print("🐛 ===== DEBUG CHECK EXISTING RESULTS =====")
        print(f"Current case ID: {self.current_case_id}")
        print(f"Output dir: {self.output_dir}")
        print(f"Analysis results: {len(self.analysis_results) if self.analysis_results else 0} items")
        
        # Kiểm tra database
        try:
            from models.db_manager import DatabaseManager
            db = DatabaseManager()
            db.connect()
            results = db.get_results_by_case(self.current_case_id)
            registry_results = [r for r in results if 'registry' in r.get('tool_used', '').lower()]
            db.disconnect()
            
            print(f"Database results: {len(results)} total, {len(registry_results)} registry")
            for r in registry_results:
                print(f"  - {r.get('tool_used', 'N/A')}: {r.get('summary', 'N/A')}")
        except Exception as e:
            print(f"Error checking database: {e}")
        
        # Kiểm tra CSV files
        if self.output_dir and os.path.exists(self.output_dir):
            try:
                import glob
                csv_files = glob.glob(os.path.join(self.output_dir, "*_analysis.csv"))
                print(f"CSV files found: {len(csv_files)}")
                for csv_file in csv_files:
                    print(f"  - {csv_file}")
            except Exception as e:
                print(f"Error checking CSV files: {e}")
        else:
            print("Output dir does not exist")
            
        # Kiểm tra has_existing_results
        has_results = self.has_existing_results()
        print(f"has_existing_results(): {has_results}")
        
        print("🐛 ===== END DEBUG =====")
        
        QMessageBox.information(self, "Debug", f"Check console output for debug info.\n\nCase: {self.current_case_id}\nHas results: {has_results}")
        
    def show_export_menu(self):
        """Hiển thị menu tùy chọn xuất dữ liệu."""
        menu = QMenu()

        # Thêm nút chạy phân tích mới
        menu.addAction("🔬 Chạy Phân Tích Mới", self.run_new_analysis)
        menu.addSeparator()
        
        menu.addAction("📄 Xuất Báo Cáo Word", self.export_html)
        menu.addSeparator()
        menu.addAction("📋 Tạo Mẫu Báo Cáo Word", self.create_word_template)

        menu.exec_(self.ui.btnExport.mapToGlobal(self.ui.btnExport.rect().bottomLeft()))

    def run_new_analysis(self):
        """Chạy phân tích mới (được gọi từ menu)."""
        if not self.loaded_hives:
            # Nếu chưa có hive, thử tự động load
            if self.current_case_id:
                try:
                    from models.db_manager import DatabaseManager
                    db = DatabaseManager()
                    db.connect()
                    case_info = db.get_case_with_investigator(self.current_case_id)
                    db.disconnect()
                    
                    if case_info and case_info.get('archive_path'):
                        # Thử tự động load registry files
                        self.auto_load_case_registry(case_info['archive_path'])
                        
                        # Nếu vẫn không có hive, thông báo
                        if not self.loaded_hives:
                            QMessageBox.warning(self, "Không có Hive", "Không tìm thấy registry hive trong case này.")
                            return
                    else:
                        QMessageBox.warning(self, "Không có Hive", "Vui lòng tải registry hive trước.")
                        return
                except Exception as e:
                    QMessageBox.warning(self, "Lỗi", f"Lỗi tải registry files: {str(e)}")
                    return
            else:
                QMessageBox.warning(self, "Không có Hive", "Vui lòng tải registry hive trước.")
                return
        
        # Chạy phân tích
        self.start_comprehensive_analysis()
        
    def start_silent_analysis(self):
        """Bắt đầu phân tích ngầm không hiển thị dialog."""
        if not self.loaded_hives:
            return
            
        # Kiểm tra thư mục output đã được thiết lập chưa
        if not self.output_dir:
            self._set_fallback_output_dir()

        # Tìm file batch
        batch_file = os.path.join(self.batch_dir, "DFIRBatch.reb")
        if not os.path.exists(batch_file):
            # Tìm bất kỳ file batch nào
            import glob
            batch_files = glob.glob(os.path.join(self.batch_dir, "*.reb"))
            if batch_files:
                batch_file = batch_files[0]
            else:
                self.update_status("Không tìm thấy file batch RECmd", "red")
                return

        self.update_status("Đang phân tích...", "yellow")
        
        # Tạo analysis thread (không hiển thị progress dialog)
        hive_files = list(self.loaded_hives.keys())
        self.current_analysis_thread = RegistryAnalysisThread(
            self.recmd_path, batch_file, hive_files, self.output_dir
        )

        # Kết nối signals
        self.current_analysis_thread.status_updated.connect(lambda s: self.update_status(f"RECmd: {s}", "yellow"))
        self.current_analysis_thread.analysis_completed.connect(self.on_silent_analysis_completed)
        self.current_analysis_thread.error_occurred.connect(self.on_analysis_error)

        # Bắt đầu phân tích ngầm
        self.current_analysis_thread.start()
        
    def on_silent_analysis_completed(self, results):
        """Xử lý hoàn thành phân tích ngầm."""
        print(f"🎯 on_silent_analysis_completed được gọi với {len(results)} hives")
        self.analysis_results = results
        total_records = sum(len(r) for r in results.values())

        self._save_analysis_results_to_database(results)
        self.update_status(f"Hoàn thành: {total_records} artifacts", "green")
        
            
    def start_comprehensive_analysis(self):
        """Bắt đầu phân tích RECmd."""
        if not self.loaded_hives:
            QMessageBox.warning(self, "Không có Hive", "Vui lòng tải registry hive trước.")
            return

        # Kiểm tra thư mục output đã được thiết lập chưa
        if not self.output_dir:
            QMessageBox.warning(self, "Chưa chọn Case", "Vui lòng chọn case trước hoặc thư mục output sẽ là tạm thời.")
            self._set_fallback_output_dir()

        # Tìm file batch
        batch_file = os.path.join(self.batch_dir, "DFIRBatch.reb")
        if not os.path.exists(batch_file):
            # Tìm bất kỳ file batch nào
            import glob
            batch_files = glob.glob(os.path.join(self.batch_dir, "*.reb"))
            if batch_files:
                batch_file = batch_files[0]
            else:
                QMessageBox.warning(self, "Không có Batch File", "Không tìm thấy file batch RECmd.")
                return

                
        # Tạo progress dialog
        progress = QProgressDialog("Đang chạy phân tích RECmd...", "Hủy", 0, 100, self)
        progress.setWindowModality(Qt.WindowModal)

        # Tạo analysis thread
        hive_files = list(self.loaded_hives.keys())
        self.current_analysis_thread = RegistryAnalysisThread(
            self.recmd_path, batch_file, hive_files, self.output_dir
        )

        # Kết nối signals
        self.current_analysis_thread.progress_updated.connect(progress.setValue)
        self.current_analysis_thread.status_updated.connect(lambda s: self.update_status(s, "yellow"))
        self.current_analysis_thread.analysis_completed.connect(self.on_analysis_completed)
        self.current_analysis_thread.error_occurred.connect(self.on_analysis_error)

        progress.canceled.connect(self.current_analysis_thread.cancel)

        # Bắt đầu phân tích
        self.current_analysis_thread.start()
        progress.exec_()
        
    def on_analysis_completed(self, results):
        """Xử lý hoàn thành phân tích."""
        print(f"🎯 on_analysis_completed được gọi với {len(results)} hives")
        self.analysis_results = results
        total_records = sum(len(r) for r in results.values())

        # Lưu kết quả vào database
        self._save_analysis_results_to_database(results)

        self.update_status(f"Phân tích hoàn thành: {total_records} bản ghi", "green")


    def _save_analysis_results_to_database(self, results):
        """Lưu kết quả phân tích vào database để xuất báo cáo với CoC."""
        print(f"🔄 Bắt đầu lưu {len(results)} hive results vào database")
        print(f"📋 Case ID: {self.current_case_id}")
        
        if not self.current_case_id:
            print("❌ Không có case được chọn, bỏ qua việc lưu database")
            return

        try:
            from models.db_manager import DatabaseManager
            db = DatabaseManager()
            db.connect()

            # Lưu kết quả cho từng hive file riêng biệt
            total_results_saved = 0
            
            for hive_file, hive_results in results.items():
                # Tạo artifact riêng cho từng hive file
                artifact_id = self._create_hive_artifact_with_hash(db, hive_file)
                
                if artifact_id:
                    # Phân tích loại hive
                    hive_type = self.detect_hive_type(hive_file)
                    hive_name = os.path.basename(hive_file)
                    
                    # Tạo tóm tắt chi tiết
                    summary = f"Phân tích {hive_type} Hive ({hive_name}): Tìm thấy {len(hive_results)} artifacts"
                    
                    # Lưu kết quả với đường dẫn CSV file thực tế
                    hive_name_clean = os.path.splitext(os.path.basename(hive_file))[0]
                    csv_result_path = os.path.join(self.output_dir, f"{hive_name_clean}_analysis.csv")
                    
                    print(f"🔄 Đang lưu result vào DB: artifact_id={artifact_id}, tool={hive_type}")
                    result_id = db.add_analysis_result(
                        artifact_id=artifact_id,
                        tool_used=f"RECmd Registry Analysis - {hive_type}",
                        summary=summary,
                        result_path=csv_result_path
                    )
                    print(f"✅ Đã lưu result_id: {result_id}")

                    # Log hoạt động chi tiết cho từng hive
                    if result_id:
                        db.log_activity(
                            case_id=self.current_case_id,
                            artefact_id=artifact_id,
                            action=f"REGISTRY_ANALYSIS_{hive_type}: Phân tích {hive_name} ({hive_type}): {len(hive_results)} artifacts. Result ID: {result_id}",
                            tool_used="RECmd"
                        )
                        total_results_saved += 1

            # Log tổng kết
            if total_results_saved > 0:
                db.log_activity(
                    case_id=self.current_case_id,
                    action=f"REGISTRY_ANALYSIS_BATCH_COMPLETED: Hoàn thành phân tích {total_results_saved} hive files với tổng {sum(len(r) for r in results.values())} artifacts",
                    tool_used="RECmd"
                )

            db.disconnect()
            print(f"Đã lưu {total_results_saved} kết quả phân tích registry vào database cho case {self.current_case_id}")

        except Exception as e:
            print(f"Lỗi lưu vào database: {e}")
            traceback.print_exc()

    def _create_hive_artifact_with_hash(self, db, hive_file):
        """Tạo artifact cho hive file cụ thể với hash để đảm bảo CoC."""
        try:
            hive_name = os.path.basename(hive_file)
            hive_type = self.detect_hive_type(hive_file)
            
            print(f"🔄 Tạo artifact cho {hive_name} ({hive_type})")
            
            # Tính hash SHA256 cho file
            file_hash = self._calculate_file_hash(hive_file)
            file_size = os.path.getsize(hive_file)
            
            # Tạo artifact với tên chi tiết
            artifact_name = f"Registry {hive_type} Hive - {hive_name}"
            
            artifact_id = db.add_artifact(
                case_id=self.current_case_id,
                name=artifact_name,
                source_path=hive_file,
                evidence_type=f"REGISTRY_{hive_type}",
                size=file_size,
                mime_type="application/octet-stream"
            )
            
            print(f"✅ Tạo artifact_id: {artifact_id}")
            
            if artifact_id and file_hash:
                # Lưu hash vào bảng Hashes
                db.add_hash(artifact_id, "SHA256", file_hash)
                
                # Log việc tạo artifact
                db.log_activity(
                    case_id=self.current_case_id,
                    artefact_id=artifact_id,
                    action=f"ARTIFACT_CREATED: {artifact_name}, SHA256: {file_hash}",
                    tool_used="Registry Analysis"
                )
            
            return artifact_id
            
        except Exception as e:
            print(f"Lỗi tạo artifact cho {hive_file}: {e}")
            return None

    def _calculate_file_hash(self, file_path):
        """Tính hash SHA256 cho file."""
        try:
            sha256_hash = hashlib.sha256()
            with open(file_path, "rb") as f:
                # Đọc file theo chunk để tiết kiệm memory
                for chunk in iter(lambda: f.read(4096), b""):
                    sha256_hash.update(chunk)
            return sha256_hash.hexdigest()
        except Exception as e:
            print(f"Lỗi tính hash cho {file_path}: {e}")
            return None

    def _get_or_create_registry_artifact(self, db):
        """Lấy artifact registry hiện có hoặc tạo mới."""
        try:
            # Kiểm tra xem artifact registry đã tồn tại cho case này chưa
            artifacts = db.get_artifacts_by_case(self.current_case_id)
            for artifact in artifacts:
                if artifact['evidence_type'] == 'REGISTRY_HIVE':
                    return artifact['artefact_id']

            # Tạo artifact mới cho kết quả phân tích registry
            registry_path = os.path.join(self.output_dir, "registry_analysis_results")
            os.makedirs(registry_path, exist_ok=True)

            artifact_id = db.add_artifact(
                case_id=self.current_case_id,
                name="Kết quả phân tích Registry",
                source_path=registry_path,
                evidence_type="REGISTRY_HIVE",
                mime_type="application/octet-stream"
            )

            return artifact_id

        except Exception as e:
            print(f"Lỗi lấy/tạo registry artifact: {e}")
            return None
        
    def on_analysis_error(self, error):
        """Xử lý lỗi phân tích."""
        self.update_status(f"Phân tích thất bại: {error}", "red")
        QMessageBox.critical(self, "Lỗi phân tích", f"Phân tích thất bại:\n{error}")

    def refresh_view(self):
        """Làm mới view hiện tại."""
        if self.loaded_hives:
            self.build_registry_tree()
            self.update_timeline()
            self.update_status("Đã làm mới view", "green")
            
    def export_html(self):
        """Xuất báo cáo Registry Analysis dưới dạng Word (.docx)."""
        if not DOCX_AVAILABLE:
            QMessageBox.warning(self, "Lỗi", "Thiếu thư viện python-docx. Cài đặt bằng: pip install python-docx")
            return

        if not self.current_case_id:
            QMessageBox.warning(self, "Lỗi", "Vui lòng chọn case trước khi xuất báo cáo.")
            return

        if not hasattr(self, 'reports_dir') or not self.reports_dir:
            QMessageBox.warning(self, "Lỗi", "Thư mục reports chưa được thiết lập.")
            return

        # Tự động tạo tên file với timestamp
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"registry_report_case_{self.current_case_id}_{timestamp}.docx"
        filepath = os.path.join(self.reports_dir, filename)

        try:
            self._generate_word_report(filepath)
            # Lưu thông tin báo cáo vào database với hash
            self._save_report_to_database(filepath)
            self.update_status("Đã xuất báo cáo Word", "green")
            QMessageBox.information(self, "Thành công", f"Báo cáo đã được xuất:\n{filepath}")
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", f"Lỗi xuất báo cáo: {str(e)}")

    def create_word_template(self):
        """Tạo mẫu báo cáo Word trống."""
        if not DOCX_AVAILABLE:
            QMessageBox.warning(self, "Lỗi", "Thiếu thư viện python-docx. Cài đặt bằng: pip install python-docx")
            return

        if not self.current_case_id:
            QMessageBox.warning(self, "Lỗi", "Vui lòng chọn case trước khi tạo mẫu báo cáo.")
            return

        if not hasattr(self, 'reports_dir') or not self.reports_dir:
            QMessageBox.warning(self, "Lỗi", "Thư mục reports chưa được thiết lập.")
            return

        # Tự động tạo tên file với timestamp
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"registry_template_case_{self.current_case_id}_{timestamp}.docx"
        filepath = os.path.join(self.reports_dir, filename)

        try:
            self._create_word_template(filepath)
            self.update_status("Đã tạo mẫu báo cáo Word", "green")
            QMessageBox.information(self, "Thành công", f"Mẫu báo cáo đã được tạo:\n{filepath}")
        except Exception as e:
            QMessageBox.critical(self, "Lỗi", f"Lỗi tạo mẫu báo cáo: {str(e)}")

    def _create_word_template(self, filepath):
        """Tạo mẫu báo cáo Word trống với cấu trúc chuẩn."""
        try:

            # Tạo document mới
            doc = Document()

            # Thiết lập style cho document
            self._setup_word_styles(doc)

            # Header với thông tin case
            self._add_template_header(doc)

            # Các section trống
            self._add_template_coc_section(doc)
            self._add_template_summary_section(doc)
            self._add_template_detailed_section(doc)
            self._add_template_activity_section(doc)
            self._add_template_signature_section(doc)

            # Lưu document
            doc.save(filepath)

        except ImportError:
            raise Exception("Thiếu thư viện python-docx. Cài đặt bằng: pip install python-docx")
        except Exception as e:
            raise Exception(f"Lỗi tạo mẫu báo cáo Word: {str(e)}")

    def _add_template_header(self, doc):
        """Thêm header mẫu với thông tin case."""
        # Tiêu đề chính
        title = doc.add_paragraph("BÁO CÁO PHÂN TÍCH REGISTRY", style='CustomHeading1')
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        subtitle = doc.add_paragraph("CHAIN OF CUSTODY - EVIDENCE INTEGRITY", style='CustomHeading2')
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph("")

        # Thông tin case mẫu
        info_table = doc.add_table(rows=6, cols=2)
        info_table.style = 'Table Grid'

        # Headers
        hdr_cells = info_table.rows[0].cells
        hdr_cells[0].text = "Thông tin Case"
        hdr_cells[1].text = "Giá trị"

        # Case ID
        row_cells = info_table.rows[1].cells
        row_cells[0].text = "Case ID:"
        row_cells[1].text = "[Nhập Case ID]"

        # Tiêu đề case
        row_cells = info_table.rows[2].cells
        row_cells[0].text = "Tiêu đề Case:"
        row_cells[1].text = "[Nhập tiêu đề case]"

        # Điều tra viên
        row_cells = info_table.rows[3].cells
        row_cells[0].text = "Điều tra viên:"
        row_cells[1].text = "[Nhập tên điều tra viên]"

        # Ngày tạo báo cáo
        row_cells = info_table.rows[4].cells
        row_cells[0].text = "Ngày tạo báo cáo:"
        row_cells[1].text = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

        # Báo cáo ID
        row_cells = info_table.rows[5].cells
        row_cells[0].text = "Báo cáo ID:"
        row_cells[1].text = f"REG_[Case ID]_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

        doc.add_page_break()

    def _add_template_coc_section(self, doc):
        """Thêm section Chain of Custody mẫu."""
        heading = doc.add_paragraph("CHAIN OF CUSTODY - EVIDENCE INTEGRITY", style='CustomHeading2')
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Tạo bảng artifacts mẫu
        table = doc.add_table(rows=2, cols=6)
        table.style = 'Table Grid'

        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Artifact"
        hdr_cells[1].text = "Loại"
        hdr_cells[2].text = "Đường dẫn"
        hdr_cells[3].text = "Kích thước"
        hdr_cells[4].text = "SHA256 Hash"
        hdr_cells[5].text = "Thời gian thu thập"

        # Sample row
        row_cells = table.rows[1].cells
        row_cells[0].text = "[Tên artifact]"
        row_cells[1].text = "[Loại evidence]"
        row_cells[2].text = "[Đường dẫn file]"
        row_cells[3].text = "[Kích thước file]"
        row_cells[4].text = "[Hash SHA256]"
        row_cells[5].text = "[Thời gian thu thập]"

        doc.add_paragraph("[Mô tả chi tiết về Chain of Custody]")
        doc.add_page_break()

    def _add_template_summary_section(self, doc):
        """Thêm section tóm tắt mẫu."""
        heading = doc.add_paragraph("TÓM TẮT PHÂN TÍCH", style='CustomHeading2')

        summary_table = doc.add_table(rows=5, cols=2)
        summary_table.style = 'Table Grid'

        # Số lượng kết quả phân tích
        row_cells = summary_table.rows[0].cells
        row_cells[0].text = "Số lượng kết quả phân tích:"
        row_cells[1].text = "[Nhập số lượng]"

        # Số lượng hoạt động đã log
        row_cells = summary_table.rows[1].cells
        row_cells[0].text = "Số lượng hoạt động đã log:"
        row_cells[1].text = "[Nhập số lượng]"

        # Số loại hive được phân tích
        row_cells = summary_table.rows[2].cells
        row_cells[0].text = "Số loại hive được phân tích:"
        row_cells[1].text = "[Nhập số loại hive]"

        # Công cụ sử dụng
        row_cells = summary_table.rows[3].cells
        row_cells[0].text = "Công cụ sử dụng:"
        row_cells[1].text = "RECmd (Registry Explorer Command Line)"

        # Tổng số artifacts
        row_cells = summary_table.rows[4].cells
        row_cells[0].text = "Tổng số artifacts:"
        row_cells[1].text = "[Nhập tổng số]"

        doc.add_paragraph("[Mô tả chi tiết về kết quả phân tích]")
        doc.add_page_break()

    def _add_template_detailed_section(self, doc):
        """Thêm section chi tiết kết quả mẫu."""
        heading = doc.add_paragraph("CHI TIẾT KẾT QUẢ PHÂN TÍCH THEO HIVE", style='CustomHeading2')

        table = doc.add_table(rows=2, cols=6)
        table.style = 'Table Grid'

        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "ID"
        hdr_cells[1].text = "Loại Hive"
        hdr_cells[2].text = "Công cụ"
        hdr_cells[3].text = "Tóm tắt"
        hdr_cells[4].text = "Đường dẫn kết quả"
        hdr_cells[5].text = "Thời gian chạy"

        # Sample row
        row_cells = table.rows[1].cells
        row_cells[0].text = "[ID kết quả]"
        row_cells[1].text = "[Loại hive: SYSTEM/SOFTWARE/SAM/ETC]"
        row_cells[2].text = "[Công cụ sử dụng]"
        row_cells[3].text = "[Tóm tắt kết quả]"
        row_cells[4].text = "[Đường dẫn file kết quả]"
        row_cells[5].text = "[Thời gian phân tích]"

        doc.add_paragraph("[Mô tả chi tiết về từng kết quả phân tích]")
        doc.add_page_break()

    def _add_template_activity_section(self, doc):
        """Thêm section lịch sử hoạt động mẫu."""
        heading = doc.add_paragraph("LỊCH SỬ HOẠT ĐỘNG (CHAIN OF CUSTODY LOG)", style='CustomHeading2')

        table = doc.add_table(rows=2, cols=4)
        table.style = 'Table Grid'

        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Thời gian"
        hdr_cells[1].text = "Hành động"
        hdr_cells[2].text = "Công cụ"
        hdr_cells[3].text = "Chi tiết"

        # Sample row
        row_cells = table.rows[1].cells
        row_cells[0].text = "[Thời gian hoạt động]"
        row_cells[1].text = "[Hành động đã thực hiện]"
        row_cells[2].text = "[Công cụ sử dụng]"
        row_cells[3].text = "[Chi tiết hoạt động]"

        doc.add_paragraph("[Mô tả chi tiết về lịch sử hoạt động]")
        doc.add_page_break()

    def _add_template_signature_section(self, doc):
        """Thêm section xác nhận điều tra viên mẫu."""
        heading = doc.add_paragraph("XÁC NHẬN ĐIỀU TRA VIÊN", style='CustomHeading2')
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph("Tôi xác nhận rằng:")

        # Danh sách xác nhận mẫu
        confirmations = [
            "Tất cả evidence đã được xử lý theo đúng quy trình Chain of Custody",
            "Hash values đã được kiểm tra để đảm bảo tính toàn vẹn dữ liệu",
            "Các công cụ phân tích được sử dụng là đáng tin cậy và được kiểm định",
            "Kết quả phân tích được lưu trữ an toàn và có thể truy xuất"
        ]

        for confirmation in confirmations:
            p = doc.add_paragraph(confirmation, style='List Bullet')

        doc.add_paragraph("")

        # Thông tin ký tên mẫu
        signature_table = doc.add_table(rows=3, cols=2)
        signature_table.style = 'Table Grid'

        # Điều tra viên
        row_cells = signature_table.rows[0].cells
        row_cells[0].text = "Điều tra viên:"
        row_cells[1].text = "[Nhập tên điều tra viên]"

        # Ngày ký
        row_cells = signature_table.rows[1].cells
        row_cells[0].text = "Ngày ký:"
        row_cells[1].text = datetime.now().strftime('%Y-%m-%d')

        # Chữ ký
        row_cells = signature_table.rows[2].cells
        row_cells[0].text = "Chữ ký:"
        row_cells[1].text = "_______________________"


    def _generate_word_report(self, filepath):
        """Tạo báo cáo Word (.docx) từ dữ liệu database."""
        try:
            # Kiểm tra thư viện python-docx
            if not DOCX_AVAILABLE:
                raise Exception("Thiếu thư viện python-docx. Cài đặt bằng: pip install python-docx")

            from models.db_manager import DatabaseManager
            db = DatabaseManager()
            db.connect()

            # Lấy thông tin case
            case_info = db.get_case_with_investigator(self.current_case_id) if self.current_case_id else None

            # Lấy kết quả phân tích
            results = db.get_results_by_case(self.current_case_id) if self.current_case_id else []
            registry_results = [r for r in results if 'registry' in r.get('tool_used', '').lower()]

            # Lấy activity logs
            activity_logs = db.get_activity_logs(case_id=self.current_case_id) if self.current_case_id else []
            registry_logs = [log for log in activity_logs if 'registry' in log.get('action', '').lower()]

            # Tạo document Word
            self._create_word_report_content(filepath, case_info, registry_results, registry_logs)

            db.disconnect()

        except Exception as e:
            raise Exception(f"Lỗi tạo báo cáo Word: {str(e)}")



    def _create_word_report_content(self, filepath, case_info, registry_results, registry_logs):
        """Tạo nội dung Word (.docx) cho báo cáo với đầy đủ thông tin CoC."""

        # Tạo document mới
        doc = Document()

        # Thiết lập style cho document
        self._setup_word_styles(doc)

        # Lấy thông tin artifacts và hash để đảm bảo CoC
        artifacts_info = self._get_artifacts_with_hash() if self.current_case_id else []

        # Header với thông tin case
        self._add_report_header(doc, case_info)

        # Chain of Custody Section
        self._add_coc_section(doc, artifacts_info)

        # Tóm tắt phân tích
        self._add_summary_section(doc, registry_results, registry_logs)

        # Chi tiết kết quả phân tích theo hive
        self._add_detailed_results_section(doc, registry_results)

        # Lịch sử hoạt động
        self._add_activity_log_section(doc, registry_logs)

        # Phần xác nhận điều tra viên
        self._add_signature_section(doc, case_info)

        # Lưu document
        doc.save(filepath)

    def _setup_word_styles(self, doc):
        """Thiết lập các style cho document Word."""

        # Style cho heading 1
        styles = doc.styles
        h1_style = styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
        h1_style.font.size = Pt(18)
        h1_style.font.bold = True
        h1_style.font.color.rgb = RGBColor(31, 73, 125)

        # Style cho heading 2
        h2_style = styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
        h2_style.font.size = Pt(14)
        h2_style.font.bold = True
        h2_style.font.color.rgb = RGBColor(79, 129, 189)

        # Style cho normal text
        normal_style = styles['Normal']
        normal_style.font.size = Pt(11)
        normal_style.font.name = 'Times New Roman'

    def _add_report_header(self, doc, case_info):
        """Thêm header báo cáo với thông tin case."""
        # Tiêu đề chính
        title = doc.add_paragraph("BÁO CÁO PHÂN TÍCH REGISTRY", style='CustomHeading1')
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        subtitle = doc.add_paragraph("CHAIN OF CUSTODY - EVIDENCE INTEGRITY", style='CustomHeading2')
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph("")

        # Thông tin case
        info_table = doc.add_table(rows=6, cols=2)
        info_table.style = 'Table Grid'

        # Headers
        hdr_cells = info_table.rows[0].cells
        hdr_cells[0].text = "Thông tin Case"
        hdr_cells[1].text = "Giá trị"

        # Case ID
        row_cells = info_table.rows[1].cells
        row_cells[0].text = "Case ID:"
        row_cells[1].text = case_info.get('case_id', 'N/A') if case_info else 'N/A'

        # Tiêu đề case
        row_cells = info_table.rows[2].cells
        row_cells[0].text = "Tiêu đề Case:"
        row_cells[1].text = case_info.get('title', 'N/A') if case_info else 'N/A'

        # Điều tra viên
        row_cells = info_table.rows[3].cells
        row_cells[0].text = "Điều tra viên:"
        row_cells[1].text = case_info.get('full_name', 'N/A') if case_info else 'N/A'

        # Ngày tạo báo cáo
        row_cells = info_table.rows[4].cells
        row_cells[0].text = "Ngày tạo báo cáo:"
        row_cells[1].text = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

        # Báo cáo ID
        row_cells = info_table.rows[5].cells
        row_cells[0].text = "Báo cáo ID:"
        row_cells[1].text = f"REG_{self.current_case_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

        doc.add_page_break()

    def _add_coc_section(self, doc, artifacts_info):
        """Thêm section Chain of Custody."""
        heading = doc.add_paragraph("CHAIN OF CUSTODY - EVIDENCE INTEGRITY", style='CustomHeading2')
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        if artifacts_info:
            # Tạo bảng artifacts
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'

            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Artifact"
            hdr_cells[1].text = "Loại"
            hdr_cells[2].text = "Đường dẫn"
            hdr_cells[3].text = "Kích thước"
            hdr_cells[4].text = "SHA256 Hash"
            hdr_cells[5].text = "Thời gian thu thập"

            # Data rows
            for artifact in artifacts_info:
                row_cells = table.add_row().cells
                row_cells[0].text = artifact.get('name', 'N/A')
                row_cells[1].text = artifact.get('evidence_type', 'N/A')
                row_cells[2].text = artifact.get('source_path', 'N/A')
                row_cells[3].text = f"{artifact.get('size', 0):,} bytes"
                row_cells[4].text = artifact.get('sha256', 'N/A')
                row_cells[5].text = artifact.get('collected_at', 'N/A')

            doc.add_paragraph(f"Tổng số artifacts được phân tích: {len(artifacts_info)}")
        else:
            doc.add_paragraph("Không có thông tin artifacts.")

        doc.add_page_break()

    def _add_summary_section(self, doc, registry_results, registry_logs):
        """Thêm section tóm tắt phân tích."""
        heading = doc.add_paragraph("TÓM TẮT PHÂN TÍCH", style='CustomHeading2')

        summary_table = doc.add_table(rows=5, cols=2)
        summary_table.style = 'Table Grid'

        # Số lượng kết quả phân tích
        row_cells = summary_table.rows[0].cells
        row_cells[0].text = "Số lượng kết quả phân tích:"
        row_cells[1].text = str(len(registry_results))

        # Số lượng hoạt động đã log
        row_cells = summary_table.rows[1].cells
        row_cells[0].text = "Số lượng hoạt động đã log:"
        row_cells[1].text = str(len(registry_logs))

        # Số loại hive được phân tích
        hive_types = set()
        for r in registry_results:
            tool_used = r.get('tool_used', '')
            if ' - ' in tool_used:
                hive_types.add(tool_used.split(' - ')[-1])

        row_cells = summary_table.rows[2].cells
        row_cells[0].text = "Số loại hive được phân tích:"
        row_cells[1].text = str(len(hive_types))

        # Công cụ sử dụng
        row_cells = summary_table.rows[3].cells
        row_cells[0].text = "Công cụ sử dụng:"
        row_cells[1].text = "RECmd (Registry Explorer Command Line)"

        # Tổng số artifacts
        total_artifacts = sum(len(r) for r in self.analysis_results.values()) if self.analysis_results else 0
        row_cells = summary_table.rows[4].cells
        row_cells[0].text = "Tổng số artifacts:"
        row_cells[1].text = str(total_artifacts)

        doc.add_page_break()

    def _add_detailed_results_section(self, doc, registry_results):
        """Thêm section chi tiết kết quả phân tích theo hive."""
        heading = doc.add_paragraph("CHI TIẾT KẾT QUẢ PHÂN TÍCH THEO HIVE", style='CustomHeading2')

        if registry_results:
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'

            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "ID"
            hdr_cells[1].text = "Loại Hive"
            hdr_cells[2].text = "Công cụ"
            hdr_cells[3].text = "Tóm tắt"
            hdr_cells[4].text = "Đường dẫn kết quả"
            hdr_cells[5].text = "Thời gian chạy"

            # Data rows
            for result in registry_results:
                row_cells = table.add_row().cells
                hive_type = result.get('tool_used', '').split(' - ')[-1] if ' - ' in result.get('tool_used', '') else 'Unknown'

                row_cells[0].text = str(result.get('result_id', 'N/A'))
                row_cells[1].text = hive_type
                row_cells[2].text = result.get('tool_used', 'N/A')
                row_cells[3].text = result.get('summary', 'N/A')
                row_cells[4].text = result.get('result_path', 'N/A')
                row_cells[5].text = result.get('run_at', 'N/A')
        else:
            doc.add_paragraph("Không có kết quả phân tích nào.")

        doc.add_page_break()

    def _add_activity_log_section(self, doc, registry_logs):
        """Thêm section lịch sử hoạt động."""
        heading = doc.add_paragraph("LỊCH SỬ HOẠT ĐỘNG (CHAIN OF CUSTODY LOG)", style='CustomHeading2')

        if registry_logs:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'

            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Thời gian"
            hdr_cells[1].text = "Hành động"
            hdr_cells[2].text = "Công cụ"
            hdr_cells[3].text = "Chi tiết"

            # Data rows
            for log in registry_logs:
                row_cells = table.add_row().cells
                row_cells[0].text = str(log.get('timestamp', 'N/A'))
                row_cells[1].text = log.get('action', 'N/A')
                row_cells[2].text = log.get('tool_used', 'N/A')
                row_cells[3].text = log.get('details', 'N/A')
        else:
            doc.add_paragraph("Không có lịch sử hoạt động nào.")

        doc.add_page_break()

    def _add_signature_section(self, doc, case_info):
        """Thêm section xác nhận điều tra viên."""
        heading = doc.add_paragraph("XÁC NHẬN ĐIỀU TRA VIÊN", style='CustomHeading2')
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph("Tôi xác nhận rằng:")

        # Danh sách xác nhận
        confirmations = [
            "Tất cả evidence đã được xử lý theo đúng quy trình Chain of Custody",
            "Hash values đã được kiểm tra để đảm bảo tính toàn vẹn dữ liệu",
            "Các công cụ phân tích được sử dụng là đáng tin cậy và được kiểm định",
            "Kết quả phân tích được lưu trữ an toàn và có thể truy xuất"
        ]

        for confirmation in confirmations:
            p = doc.add_paragraph(confirmation, style='List Bullet')

        doc.add_paragraph("")

        # Thông tin ký tên
        signature_table = doc.add_table(rows=3, cols=2)
        signature_table.style = 'Table Grid'

        # Điều tra viên
        row_cells = signature_table.rows[0].cells
        row_cells[0].text = "Điều tra viên:"
        row_cells[1].text = case_info.get('full_name', 'N/A') if case_info else 'N/A'

        # Ngày ký
        row_cells = signature_table.rows[1].cells
        row_cells[0].text = "Ngày ký:"
        row_cells[1].text = datetime.now().strftime('%Y-%m-%d')

        # Chữ ký
        row_cells = signature_table.rows[2].cells
        row_cells[0].text = "Chữ ký:"
        row_cells[1].text = "_______________________"

    def _get_artifacts_with_hash(self):
        """Lấy thông tin artifacts cùng với hash để đảm bảo CoC."""
        try:
            from models.db_manager import DatabaseManager
            db = DatabaseManager()
            db.connect()

            # Lấy artifacts registry của case hiện tại
            artifacts = db.get_artifacts_by_case(self.current_case_id)
            registry_artifacts = [a for a in artifacts if 'REGISTRY' in a.get('evidence_type', '')]

            # Lấy hash cho từng artifact
            artifacts_with_hash = []
            for artifact in registry_artifacts:
                artifact_id = artifact['artefact_id']
                hashes = self._get_hashes_by_artifact(db, artifact_id)
                
                # Lấy SHA256 hash
                sha256_hash = 'N/A'
                for hash_info in hashes:
                    if hash_info['hash_type'] == 'SHA256':
                        sha256_hash = hash_info['sha256']
                        break

                artifact_info = {
                    'name': artifact['name'],
                    'evidence_type': artifact['evidence_type'],
                    'source_path': artifact['source_path'],
                    'size': artifact['size'],
                    'collected_at': artifact['collected_at'],
                    'sha256': sha256_hash
                }
                artifacts_with_hash.append(artifact_info)

            db.disconnect()
            return artifacts_with_hash

        except Exception as e:
            print(f"Lỗi lấy thông tin artifacts: {e}")
            return []

    def _get_hashes_by_artifact(self, db, artifact_id):
        """Lấy tất cả hash của một artifact."""
        try:
            query = "SELECT * FROM Hashes WHERE artefact_id = ?"
            return db.fetch_all(query, (artifact_id,))
        except Exception as e:
            print(f"Lỗi lấy hash cho artifact {artifact_id}: {e}")
            return []

    def _save_report_to_database(self, filepath, format_type="DOCX"):
        """Lưu thông tin báo cáo Word vào database với hash để đảm bảo CoC."""
        try:
            from models.db_manager import DatabaseManager

            # Tính hash cho file báo cáo
            report_hash = self._calculate_file_hash(filepath)

            db = DatabaseManager()
            db.connect()

            # Lưu vào bảng Reports
            report_id = db.create_report(
                case_id=self.current_case_id,
                file_path=filepath,
                format="DOCX",
                sha256=report_hash
            )

            if report_id:
                # Log hoạt động
                db.log_activity(
                    case_id=self.current_case_id,
                    action="REPORT_GENERATED",
                    tool_used="Registry Analysis - DOCX",
                    details=f"Tạo báo cáo Word: {os.path.basename(filepath)}, SHA256: {report_hash}"
                )
                print(f"Đã lưu thông tin báo cáo vào database: Report ID {report_id}")

            db.disconnect()

        except Exception as e:
            print(f"Lỗi lưu báo cáo vào database: {e}")
